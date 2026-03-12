from celery import shared_task
import datetime
import os
import multiprocessing
import concurrent
import json
import PyPDF2
import fitz  # PyMuPDF
from langchain.chains import RetrievalQA
import ast
import google.generativeai as genai
import pymongo
import re
from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
import time
from django.conf import settings
from bson import ObjectId
from kafka import KafkaProducer
import urllib.parse
import logging
import requests
import os
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
import PyPDF2
from langchain.chains import RetrievalQA
from langchain_google_genai import GoogleGenerativeAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import Document
from docx import Document as Docx_document
from langchain.docstore.document import Document as LangchainDocument  # Alias langchain's Document
from langchain.chains.combine_documents import create_stuff_documents_chain
import json
import google.generativeai as genai
from phi.model.google import Gemini
from phi.agent import Agent, RunResponse
import re 
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
import PyPDF2
from langchain.chains import RetrievalQA
from langchain_google_genai import GoogleGenerativeAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
  # Alias langchain's Document
from langchain.chains.combine_documents import create_stuff_documents_chain
from phi.agent import Agent
from phi.knowledge.langchain import LangChainKnowledgeBase
from phi.model.google import Gemini
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from phi.agent import Agent, RunResponse
import os
import tempfile
import requests
from moviepy.editor import VideoFileClip, concatenate_videoclips
import io
import time
import shutil
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from moviepy.editor import VideoFileClip, concatenate_videoclips
logging.basicConfig(level=logging.INFO) 

logger = logging.getLogger(__name__)
# Load API keys from environment variables
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Set environment variables for libraries
if GOOGLE_API_KEY:
    os.environ['GOOGLE_API_KEY'] = GOOGLE_API_KEY
if GROQ_API_KEY:
    os.environ['GROQ_API_KEY'] = GROQ_API_KEY
SPRING_BOOT_URL = os.getenv("SPRING_BOOT_URL", "https://aippoint.ai/aippoint-spring-service/recording/")
PROCTORING_API_BASE_URL = "https://aippoint.ai/aippoint-django-proctoring/"

# llm = GooglePalm(google_api_key=api_key, temperature=0.7)
llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash")



# MongoDB credentials from environment variables
mongo_username = os.getenv("MONGO_USERNAME", "admin")
mongo_password = os.getenv("MONGO_PASSWORD", "")
if not mongo_password:
    password = os.getenv("MONGO_PASSWORD_PLAIN", "")
    mongo_password = urllib.parse.quote(password, safe='')
local_directory = settings.MEDIA_ROOT+'/uploads_temp'
local_directory_text = settings.MEDIA_ROOT+'/uploads_temp_txt'

local_directory_jd = settings.MEDIA_ROOT+'/uploads_jd_temp'
local_directory_text_jd = settings.MEDIA_ROOT+'/uploads_jd_temp_txt'
people_list = []
# from langchain.embeddings import GooglePalmEmbeddings
instructor_embeddings=GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
 

def save_text_file(text, text_file_path):
    try:
        if not os.path.exists(text_file_path):
            os.makedirs(os.path.dirname(text_file_path), exist_ok=True)
        
        with open(text_file_path, "w", encoding="utf-8") as txt_file:
            
            txt_file.write(text)
    except Exception as e:
        print(f'Error: {e}')
        


def clean_text(text):
    cleaned_text = re.sub(r'\s+', ' ', text)
    cleaned_text = cleaned_text.replace('/', '').replace('\\', '').replace('\n', '').replace('\"', '"').replace('\\"', '"')
    try:
        if cleaned_text.startswith("["):
            if cleaned_text.endswith("}"):
                cleaned_text = cleaned_text[:len(cleaned_text)-1]
            try:
                cleaned_text = json.loads(cleaned_text)
            except:
                cleaned_text = json.loads(cleaned_text.replace("'","\""))
        else:
            index = cleaned_text.find(":")
            cleaned_text = cleaned_text[index+1:]
            cleaned_text = cleaned_text.split(",")
    except:
        print("Exception came while cleaning")
    return cleaned_text



Instruction = [""" The Output format must follow the below format
              {
                "job_role": "",
                "job_title": "",
                "job_type": "",
                "location": "",
                "experience_required": Always a integers,
                "primary_skills": [],
                "secondary_skills": [],
                "jd_category": ""
            }
 
            for sample iam giving this output sample
            {
    "job_role": "Java Developer",
    "job_title": "Java Developer",
    "job_type": "Full-Time",
    "location": "Bengaluru, India",
    "experience_required": 1,
    "primary_skills": [
        "Java",
        "Spring",
        "Hibernate",
        "RESTful APIs",
        "SOAP",
        "MySQL",
        "PostgreSQL",
        "Oracle",
        "Git",
        "Object-oriented programming (OOP)"
    ],
    "secondary_skills": [
        "HTML",
        "CSS",
        "JavaScript",
        "Agile",
        "Microservices",
        "AWS",
        "Azure"
    ],
    "jd_category": "Java Developer"
}
 
                """]



# Extraction function
# def extract_details(job_desc):
#     details = {}
#     patterns = {
#         'job_role': r"job_role\s*:\s*(.+)",
#         'job_title': r"job_title\s*:\s*(.+)",
#         'job_type': r"job_type\s*:\s*(.+)",
#         'location': r"location\s*:\s*(.+)",
#         'experience_required': r"experience_required\s*:\s*(.+)",
#         'primary_skills': r"primary_skills\s*:\s*(\[[^\]]+\])",
#         'secondary_skills': r"secondary_skills\s*:\s*(\[[^\]]+\])",
#         'domain_specific': r"domain_specific\s*:\s*(.+)",
#         'jd_category': r"jd_category\s*:\s*(.+)"
#     }

import re
import json

def extract_details(job_desc_json):
    try:
        # Log the input
        print(f"Input JSON: {job_desc_json}")
        
        # Parse JSON string to a dictionary if it's a string
        if isinstance(job_desc_json, str) and job_desc_json.strip():
            job_desc_json = json.loads(job_desc_json)
        elif not isinstance(job_desc_json, dict):
            raise ValueError("Invalid input: Expected a JSON string or dictionary.")
        
        # Extract details directly from the dictionary
        details = {
            "job_role": job_desc_json.get("job_role", "Not Provided"),
            "job_title": job_desc_json.get("job_title", "Not Provided"),
            "job_type": job_desc_json.get("job_type", "Not Provided"),
            "mode_of_work": job_desc_json.get("mode_of_work", "Not Provided"),
            "experience_required": int(job_desc_json.get("experience_required", 0)),  # Ensure integer conversion
            "primary_skills": job_desc_json.get("primary_skills", []),
            "secondary_skills": job_desc_json.get("secondary_skills", []),
            "domain_specific": job_desc_json.get("domain_specific", []),
            "jd_category": job_desc_json.get("jd_category", "Not Provided"),
            "location": {
                "city": job_desc_json.get("location", {}).get("city", "Not Provided"),
                "state": job_desc_json.get("location", {}).get("state", "Not Provided"),
                "country": job_desc_json.get("location", {}).get("country", "Not Provided")
            }
        }

        return details

    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {str(e)}")
        return {}
    except Exception as e:
        print(f"Error extracting details: {str(e)}")
        return {}



    
    # for key, pattern in patterns.items():
    #     match = re.search(pattern, job_desc)
    #     if match:
    #         details[key] = eval(match.group(1).strip()) if 'skills' in key else match.group(1).strip()
    # return details


def process_job_description(ai_jd_data5):
    url = 'https://aippoint.ai/aippoint-spring-category/'
    try:
                
        response = requests.get(url, verify=False)
           
                # Check if the request was successful (status code 200)
        if response.status_code == 200:
                    # Parse the JSON response
            data = response.json()
               
                    # Extracting the 'category' values and storing them in 'Categories'
            Categories = [item['category'] for item in data]                  
        else:
            print(f"Failed to retrieve data. Status code: {response.status_code}")
    except Exception as e:
        print(f"An error occurred: {e}")  # Replace with the correct path if needed
           
        
        # Use the agent to query the knowledge base
    agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                instructions = Instruction,
                #   knowledge_base=knowledge_base, 
                add_references_to_prompt=True)

    # agent.print_response("give me the name from the uploaded file")
    run: RunResponse = agent.run(f"""Extract the following details from the {ai_jd_data5}. Ensure they match the instruction format:
                              "job_role": "Extract the job role from the given information after analyzing the details provided.",
                              "job_title": "Identify and extract the job title from the provided details.",
                              "job_type": "Extract the type of job (e.g., full-time, part-time, contract, remote) from the provided information.",
                              "mode_of_work": "Extract the working mode (e.g., Onsite, Hybrid, Remote) from the provided details.",
                              "location":
                                  "city": "Extract the city where the job is located. If not available, return 'Not Provided'.",
                                  "state": "Extract the state where the job is located. If not available, return 'Not Provided'.",
                                  "country": "Extract the country where the job is located. If not available, return 'Not Provided'."
                              ,
                              "experience_required": "Extract the years of experience or experience level required as indicated in the provided information. Strictly remember you must give this in an integer always.",
                              "primary_skills": "Identify and list the primary skills required for the job based on the information provided in the job description. Ensure these are formatted as a list.",
                              "secondary_skills": "Extract any secondary or additional skills mentioned in the information. Ensure these are formatted as a list.",
                              "domain_specific": "Extract the domain-specific requirements mentioned in the job description. Focus on identifying industries (e.g., healthcare, e-commerce, finance), platforms, or products explicitly stated in the JD as desired or preferred experience for the candidate. Ensure these are formatted as a list.",
                              "jd_category": "From the provided job description (JD), extract the job title or role mentioned, analyze the entire JD, including the responsibilities, skills, and qualifications, and categorize the JD based on this information. Match the extracted information with the {Categories}. If any of the details are not found from the JD, return as 'Not Provided'."
                                                        Do not include any explanations or introductory lines. The output format contains only the category. Print everything in regular text (not bold or italic)"
                            you must always generate the  output in the string format only""")
    jd_extracted_data = run.content
    print(jd_extracted_data)
    jd_cleaned_data = jd_extracted_data.strip().strip("```json").strip()
 
    print(jd_cleaned_data)
    
 
    return extract_details(jd_cleaned_data)


def pdf_to_text(blob_name):
    file_path = os.path.join(local_directory, blob_name)
    text = ""
    pdf_page_content = []  # List to hold text content of each page
 
    try:
        # Open the PDF with pdfplumber
        with pdfplumber.open(file_path) as pdf:
            # Iterate over each page and extract text
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:  # Check if the page contains any text
                    text += page_text
                    pdf_page_content.append(page_text)  # Append page text to the list
 
    except Exception as e:
        print(f'Error: {e}')
 
    file_name = os.path.splitext(blob_name)[0]
    txt_file_path = os.path.join(local_directory_text, file_name + ".txt")
 
    print("Text file path:", txt_file_path)
    print("PDF file path:", file_path)
 
    save_text_file(text, txt_file_path)
    print("Text extraction and saving completed")
 
    data = {
        "full_text": text,      
        "pdf_page_content": pdf_page_content  
    }
 
    os.remove(file_path)
    os.remove(txt_file_path)
 
    return data

def docx_to_text(blob_name):
    file_path = os.path.join(local_directory, blob_name)
    doc_page_content = []  # List to hold text content of each paragraph
 
    try:
        # Ensure the file exists and is a valid docx file
        if not os.path.isfile(file_path) or not file_path.endswith('.docx'):
            raise ValueError("Invalid file or file format. Only .docx files are supported.")
       
        # Open the document using python-docx
        doc = Docx_document(file_path)
       
        # Iterate through paragraphs and collect text
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()  # Strip unnecessary whitespace
            if paragraph_text:  # Check if the paragraph contains text
                doc_page_content.append(paragraph_text)  # Append paragraph text to list
       
        # Join all paragraphs to create the full text
        text = "\n".join(doc_page_content)
       
        # Save the extracted text to a .txt file
        file_name = os.path.splitext(blob_name)[0]
        txt_file_path = os.path.join(local_directory_text, file_name + ".txt")
        with open(txt_file_path, 'w', encoding='utf-8') as f:
            f.write(text)
 
       
        data = {
            "full_text": text.strip(),  # Full text extracted from the document
            "doc_page_content": doc_page_content  # Text content of each paragraph
        }
        print("Document data constructed:", data)
        return data
 
    except Exception as e:
        print(f'Error reading DOCX file: {e}')
        return None



 
#             # Create Document objects from page content for embedding
#             documents = [Document(page_content=text) for text in page_content]
#             print("this us the guyuhikkiu",documents)
                        
#             url = 'http://localhost:3000/parseez-spring-category'  # Replace with the correct path if needed
#             url = 'https://parseez.com/parseez-spring-category'  # Replace with the correct path if needed
            
#             try:
#                 # Making a GET request to the API
#                 response = requests.get(url)
            
#                 # Check if the request was successful (status code 200)
#                 if response.status_code == 200:
#                     # Parse the JSON response
#                     data = response.json()
                
#                     # Extracting the 'category' values and storing them in 'Categories'
#                     Categories = [item['category'] for item in data]
                
#                     print("Categories:", Categories)
#                 else:
#                     print(f"Failed to retrieve data. Status code: {response.status_code}")
#             except Exception as e:
#                 print(f"An error occurred: {e}")

            
            
#             try:
#                 instructor_embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
#                 print("Google GenerativeAi embeddings initialized successfully.")
#             except Exception as e:
#                 print(f"Error initializing Google GenerativeAi embeddings: {e}")
#                 # Optionally, log the traceback for deeper debugging insights
#                 import traceback
#                 print(traceback.format_exc())
#                 return

#             try:
#                 # Embed the documents using FAISS and instructor embeddings
#                 vectordb = FAISS.from_documents(documents=documents, embedding=instructor_embeddings)
#                 retriever = vectordb.as_retriever(score_threshold=0.7)
#                 print("FAISS database and retriever setup completed.")

#             except Exception as e:
#                 print(f"Error setting up FAISS database: {e}")
#                 return

#             # Setting up the Retrieval QA chain with the retriever
           
#             chain = RetrievalQA.from_chain_type(llm=llm,
#                                                 chain_type="stuff",
#                                                 retriever=retriever,
#                                                 input_key="query",
#                                                 return_source_documents=True)
            
#             constraint ="""You are an AI bot designed to act as a professional for parsed resume content. 
#             You are given with resume parsed content and your job is to extract the following information from the resume"""
            
#             question_1="give only name of candidate in less than 20 words, give me only name not else?"
#             name = chain.invoke(constraint + question_1)['result']
#             response = chain.invoke(constraint + question_1)
#             print(f"Response for question 1: {response}")
#             print(f"Type for question 1: ",type(response))
#             name=response['result']
   

#             print("thissssss",name)
#             # name = response['result']  # Ensure this is actually a dictionary with 'result' key

#             question_2="give me phone numbers not else and if not available give output as missing"
#             # phone_no = chain.invoke(constraint + question_2)['result']
#             response = chain.invoke(constraint + question_2)
#             phone_no=response['result']
#             print("thissssss is phone_no",phone_no)

#             question_3="give me only an email not else, if information not available give output as missing"
#             # email = chain.invoke(constraint + question_3)['result']
#             response = chain.invoke(constraint + question_3)
#             email=response['result']
#             print("thissssss is email",email)

#             question_4='in which city candidate live, only give the city name not else, if information not available give output as missing'
#             # location = chain.invoke(constraint + question_4)['result']
#             response = chain.invoke(constraint + question_4)
#             location=response['result']
#             print("thissssss is location",location)

#             question_5='give me all the list of skills and tools separated by a comma and store them in a list'
#             # skills = chain.invoke(constraint + question_5)['result']
#             # skills = chain.invoke(constraint + question_5)['result']
#             response = chain.invoke(constraint + question_5)
#             skills=response['result']
#             print("thissssss is skills",skills)


#             question_6='give me overall experience of the candidate in year and give answer in number only, if experiance information not mentioned then give output as missing'
#             # experiance_in_number = chain.invoke(constraint + question_6)['result']
#             response = chain.invoke(constraint + question_6)
#             experiance_in_number=response['result']
#             print("thissssss is experiance_in_number",experiance_in_number)

#             question_7='give me overall experience of the candidate in year and month if only mentioned by the candidate, if not mentioned by the candidate then give output as missing'
#             # experiance = chain.invoke(constraint + question_7)['result']
#             response = chain.invoke(constraint + question_7)
#             experiance=response['result']
#             print("thissssss is experiance",experiance)

#             question_8='give me all URLs of the candidate separated by a comma, if information not available give output as missing'
#             # url = chain.invoke(constraint + question_8)['result']
#             response = chain.invoke(constraint + question_8)
#             url=response['result']
#             print("thissssss is url",url)

#             question_9='give me information in the dictionary format like {college_name: degree: marks, if marks not available give output as missing} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # education_info = chain.invoke(constraint + question_9)['result']
#             response = chain.invoke(constraint + question_9)
#             education_info=response['result']
#             print("thissssss is education_info",education_info)
            
#             question_10='give the list of certification names and separate them by a comma and append in a list, if not available give output as missing'
#             # certification = chain.invoke(constraint + question_10)['result']
#             response = chain.invoke(constraint + question_10)
#             certification=response['result']
#             print("thissssss is certification",certification)



#             question_11='give me information in the dictionary format like {project_name: description_of_project} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # project_info = chain.invoke(constraint + question_11)['result']
#             response = chain.invoke(constraint + question_11)
#             project_info=response['result']
#             print("thissssss is project_info",project_info)


#             question_12='give me information only in the dictionary format like {company_name: designation_at_company: start_date_in_year_month: end_date_in_year_month} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # work = chain.invoke(constraint + question_12)['result']
#             response = chain.invoke(constraint + question_12)
#             work=response['result']
#             print("thissssss is work",work)

#             question_13=f"""From the provided resume data, extract the job title or job role that the resume is associated with,
#                       analyze the entire resume, including the skills and projects, and categorize the resume based on this information
#                       and match with the {Categories}.
#                       Do not include any explanations or introductory lines. The output format contains only the category. Print everything in regular text (not bold or italic)."""
#             Resume_Category = chain.invoke(constraint + question_13)['result']
#             print("thissssss is Resume_Category",Resume_Category)
            


#             dic = {
#                 "name": name,
#                 "email": email,
#                 "phone": phone_no,
#                 'experience': experiance,
#                 "experiance_in_number":experiance_in_number,
#                 "url": url,
#                 "location": location,
#                 "work": clean_text(work), 
#                 "education": clean_text(education_info),
#                 "certificates": clean_text(certification),
#                 "skills": clean_text(skills), 
#                 "projects": clean_text(project_info),
#                 "created_by": created_by,
#                 "req_no": req_no,
#                 "Resume_Category":Resume_Category
#             }
#             dic['file_name'] = blob_name
#             dic['Resume_Category']= Resume_Category
#             print("dicccccccccccccccc",dic)
#             print("dicccccccccccccccc",dic['file_name'])
#             client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
#             db = client["resume_parser"]
#             collection = db["resume_collection"]
#             existing_document = collection.find_one({"email": {"$regex": re.compile('^' + re.escape(dic['email']) + '$', re.IGNORECASE)}})
#             print("dattttttaaaaaaaaaaa",existing_document)
#             if existing_document is None:
#                 # If the document doesn't exist, insert the new data into the collection
#                 collection.insert_one(dic)
#                 inserted_id = collection.inserted_id
#                 file_data = {
#                     "file_name": blob_name,
#                     "req_no": req_no,
#                     "uploaded_by": created_by,
#                     "original_file_name": original_file_name,
#                     "duplicate": "no",
#                     "Resume_Category":Resume_Category
#                 }
#                 uploaded_files = db["uploaded_files"]
#                 # Insert the JSON data into the collection
#                 uploaded_files.insert_one(file_data)

                
#                 try:
#                     # Process resume data
#                     resume_skills = dic["skills"]
#                     resume_exp = dic["experiance_in_number"]
#                     created_by = dic["created_by"]
                    
#                     client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
#                     db = client['resume_parser']
#                     score_resume_by_new_jd = db['score_resume_by_new_jd']
#                     jd_collection = db['jd_collecction']

#                     jds = list(jd_collection.find())

#                     for jd in jds:
#                         try:
#                             JD_skills = jd.get('skills', [])
#                             JD_exp = jd.get('exp', 0.0)
#                             jd_id = jd.get('jd_id', '')

#                             # Calculate score for the resume against the current JD
#                             score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)

#                             scored_resume = {
#                                 'resume_id': str(inserted_id),
#                                 'resume_data': dic,
#                                 'score': score,
#                                 'jd_data': jd,
#                                 'created_by': created_by
#                             }

#                             query = {"jd_data.jd_id": jd_id, "resume_id": str(inserted_id)}
#                             # Attempt to find the document
#                             existing_document = score_resume_by_new_jd.find_one(query)

#                             if existing_document:
#                                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
#                                 score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
#                             else:
#                                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
#                                 score_resume_by_new_jd.insert_one(scored_resume)
#                         except Exception as e:
#                             logger.error(f'Error scoring resume against JDs: {str(e)}')
#                     logger.info('Scored resume against all JDs and saved successfully.')
#                 except Exception as e:
#                     logger.error(f'Error scoring resume against JDs: {str(e)}')
#             else:
#                 # Specify the collection name (replace 'mycollection' with your collection name)
#                 file_data = {
#                     "file_name": blob_name,
#                     "req_no": req_no,
#                     "uploaded_by": created_by,
#                     "original_file_name": original_file_name,
#                     "duplicate": "yes"
#                 }
#                 uploaded_files = db["uploaded_files"]
#                 # Insert the JSON data into the collection
#                 uploaded_files.insert_one(file_data)
#             client.close()

            
# #         elif blob_name.endswith(".docx"):
# #             doc_data = docx_to_text(blob_name)
# #             page_content = doc_data.get("page_content", [])
# #             print("This is page_contentttttttttttttttt",page_content)
# #             # Ensure the extracted page content is not empty
# #             if not page_content:
# #                 print("No content extracted from the docx.")
# #                 return

# #             # doc_documents = [text for text in page_contentdoc]
# #             # documents = [Document(text=text) for text in page_content]
# #             # Assuming the Document class is from a library like langchain and requires a 'text' argument
# #             documents = [Document(text=text, metadata={"source": "page_content"}) for text in page_content] # Replace 'text' with the correct argument if needed

# #             # Treat the page content directly as plain text for FAISS
# #   # Treat the page content directly as plain text for FAISS
# #             print("Documents prepared for FAISS:", documents)
# #             print("this us the guyuhikkiu",documents)

# #             print("this us the docdata",documents)
#         elif blob_name.endswith(".docx"):
#             doc_data = docx_to_text(blob_name)
#             if doc_data is None:
#                 print("Error extracting content from the docx file.")
#                 return

#             page_content = doc_data.get("page_content", [])
#             print("This is page_content:", page_content)

#             # Ensure the extracted page content is not empty
#             if not page_content:
#                 print("No content extracted from the docx.")
#                 return

#             # Use langchain's Document class to process the text for FAISS or other use cases
#             documents = [LangchainDocument(page_content=text, metadata={"source": "page_content"}) for text in page_content]

#             # Treat the page content directly as plain text for FAISS
#             print("Documents prepared for FAISS:", documents)
#             print("this is the documents:", documents)
#             url = 'http://localhost:3000/parseez-spring-category'  # Replace with the correct path if needed
#             url = 'https://parseez.com/parseez-spring-category'  # Replace with the correct path if needed
            
#             try:
#                 # Making a GET request to the API
#                 response = requests.get(url)
            
#                 # Check if the request was successful (status code 200)
#                 if response.status_code == 200:
#                     # Parse the JSON response
#                     data = response.json()
                
#                     # Extracting the 'category' values and storing them in 'Categories'
#                     Categories = [item['category'] for item in data]
                
                   
#                 else:
#                     print(f"Failed to retrieve data. Status code: {response.status_code}")
#             except Exception as e:
#                 print(f"An error occurred: {e}")            
#             # Create a FAISS instance for vector database from 'data'
#             # vectordb = FAISS.from_documents(documents=data, embedding=instructor_embeddings)

#             try:
#                 instructor_embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
#                 print("Google GenerativeAi embeddings initialized successfully.")
#             except Exception as e:
#                 print(f"Error initializing Google GenerativeAi embeddings: {e}")
#                 # Optionally, log the traceback for deeper debugging insights
#                 import traceback
#                 print(traceback.format_exc())
#                 return

#             try:
#                 # Embed the documents using FAISS and instructor embeddings
#                 vectordb = FAISS.from_documents(documents=documents, embedding=instructor_embeddings)
#                 retriever = vectordb.as_retriever(score_threshold=0.7)
#                 print("FAISS database and retriever setup completed.")

#             except Exception as e:
#                 print(f"Error setting up FAISS database: {e}")
#                 return
            
#             # Create a retriever for querying the vector database
#             # retriever = vectordb.as_retriever(score_threshold=0.7)

#             chain = RetrievalQA.from_chain_type(llm=llm,
#                                                 chain_type="stuff",
#                                                 retriever=retriever,
#                                                 input_key="query",
#                                                 return_source_documents=True)

#             # name = chain('give only name of candidate in less than 20 words, give me only name not else?')
#             # phone_no = chain('give me phone numbers not else and if not available give output as missing')
#             # email = chain('give me only an email not else, if information not available give output as missing')
#             # location = chain('in which city candidate live, only give the city name not else, if information not available give output as missing')
#             # skills = chain('give me all the list of skills and tools separated by a comma and store them in a list')
#             # experiance_in_number = chain('give me overall experience of the candidate in year and give answer in number only, if experiance information not mentioned then give output as missing')
#             # experiance = chain('give me overall experience of the candidate in year and month if only mentioned by the candidate, if not mentioned by the candidate then give output as missing')
#             # url = chain('give me all URLs of the candidate separated by a comma, if information not available give output as missing')
#             # education_info = chain('give me information in the dictionary format like {college_name: degree: marks, if marks not available give output as missing} separated by a comma and append all dictionaries into a list, if information no available give output as missing')
#             # certification = chain('give the list of certification names and separate them by a comma and append in a list, if not available give output as missing')
#             # project_info = chain('give me information in the dictionary format like {project_name: description_of_project} separated by a comma and append all dictionaries into a list, if information no available give output as missing')
#             # work = chain('give me information only in the dictionary format like {company_name: designation_at_company: start_date_in_year_month: end_date_in_year_month} separated by a comma and append all dictionaries into a list, if information no available give output as missing')


#             constraint ="Strictly answer the question in detailed list on the provided data only otherwise say missing. "
            
#             question_1="give me the name of candidate mentioned in the document, just give me the nothing aprt from that."
#             name = chain.invoke(constraint + question_1)['result']
#             print("thissssss is Name : ",name)
           
           
#             question_2="give me the Mobile number or Contact number(which may start from +91) in the document, just mobile number in the output nothing apart from that."
#             phone_no = chain.invoke(constraint + question_2)['result']
#             print("thissssss is phone_no",phone_no)
           
#             question_3="give me the email of the of the candidate mentioned in the document, just email in the output nothing apart from that. "
#             email = chain.invoke(constraint + question_3)['result']
#             print("thissssss is email",email)
           
#             question_4="give me the location of the candiadate or city mentioned in the document, just city name in the output nothing apart from that. if the city or location not found give as 'missing'"
#             # location = chain.invoke(constraint + question_4)['result']
#             location = chain.invoke(constraint + question_4)['result']
#             print("thissssss is location",location)
           
#             question_5='give me the list of all skills and tools separated by a comma and store them in a list, just give me the skills in the output nothing apart from that '
#             # skills = chain.invoke(constraint + question_5)['result']
#             # skills = chain.invoke(constraint + question_5)['result']
#             response = chain.invoke(constraint + question_5)
#             skills=response['result']
#             print("thissssss is skills",skills)
           
#             question_6='give me overall experience of the candidate in year and give answer in number only'
#             # experiance_in_number = chain.invoke(constraint + question_6)['result']
#             response = chain.invoke(constraint + question_6)
#             experiance_in_number=response['result']
#             print("thissssss is experiance_in_number",experiance_in_number)
           
#             question_7='give me overall experience of the candidate in year and month if only mentioned by the candidate, if not mentioned by the candidate then give output as missing'
#             # experiance = chain.invoke(constraint + question_7)['result']
#             response = chain.invoke(constraint + question_7)
#             experiance=response['result']
#             print("thissssss is experiance",experiance)
           
#             question_8='give me the  URLs mentioned in the resume data of the candidate separated by a comma, if information not available give output as missing'
#             # url = chain.invoke(constraint + question_8)['result']
#             response = chain.invoke(constraint + question_8)
#             url=response['result']
#             print("thissssss is url",url)
           
#             question_9='give me the education information in the dictionary format like {college_name: degree: marks, if marks not available give output as missing} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # education_info = chain.invoke(constraint + question_9)['result']
#             response = chain.invoke(constraint + question_9)
#             education_info=response['result']
#             print("thissssss is education_info",education_info)
           
#             question_10='give the list of certification names and separate them by a comma and append in a list, if not available give output as missing'
#             # certification = chain.invoke(constraint + question_10)['result']
#             response = chain.invoke(constraint + question_10)
#             certification=response['result']
#             print("thissssss is certification",certification)
           
#             question_11='give me the projects information in the dictionary format like {project_name: description_of_project} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # project_info = chain.invoke(constraint + question_11)['result']
#             response = chain.invoke(constraint + question_11)
#             project_info=response['result']
#             print("thissssss is project_info",project_info)
           
#             question_12='give me work experience information only in the dictionary format like {company_name: designation_at_company: start_date_in_year_month: end_date_in_year_month} separated by a comma and append all dictionaries into a list, if information no available give output as missing'
#             # work = chain.invoke(constraint + question_12)['result']
#             response = chain.invoke(constraint + question_12)
#             work=response['result']
#             print("thissssss is work",work)

#             question_13=f"""From the provided resume data, extract the job title or job role that the resume is associated with,
#                       analyze the entire resume, including the skills and projects, and categorize the resume based on this information
#                       and match with the {Categories}.
#                       Do not include any explanations or introductory lines. The output format contains only the category. Print everything in regular text (not bold or italic)."""
#             Resume_Category = chain.invoke(constraint + question_13)['result']
#             print("thissssss is Resume_Category",Resume_Category)
            
#             dic = {
#                 "name": name,
#                 "email": email,
#                 "phone": phone_no,
#                 'experience': experiance,
#                 "experiance_in_number":experiance_in_number,
#                 "url": url,
#                 "location": location,
#                 "work": clean_text(work), 
#                 "education": clean_text(education_info),
#                 "certificates": clean_text(certification),
#                 "skills": clean_text(skills), 
#                 "projects": clean_text(project_info),
#                 "created_by": created_by,
#                 "req_no": req_no,
#                 "Resume_Category":Resume_Category
#             }
#             dic['file_name'] = blob_name
#             dic['Resume_Category']= Resume_Category
#             # Connect to the MongoDB server (adjust the host and port as needed)
#             client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
#             # Select a database (replace 'mydatabase' with your database name)
#             db = client["resume_parser"]
#             collection = db["resume_collection"]
#             # Close the MongoDB connection
#             # Check if similar data already exists in the collection
#             # Adjust the query based on your unique criteria
#             existing_document = collection.find_one({"email": dic['email']})
#             if existing_document is None:
#                 # If the document doesn't exist, insert the new data into the collection
#                 collection.insert_one(dic)

#                 inserted_id = collection.inserted_id
#                 # Specify the collection name (replace 'mycollection' with your collection name)
#                 file_data = {
#                     "file_name": blob_name,
#                     "req_no": req_no,
#                     "uploaded_by": created_by,
#                     "original_file_name": original_file_name,
#                     "duplicate": "no",
#                     "Resume_Category":Resume_Category
#                 }
#                 uploaded_files = db["uploaded_files"]
#                 # Insert the JSON data into the collection
#                 uploaded_files.insert_one(file_data)

#                 try:
#                     # Process resume data
#                     resume_skills = dic["skills"]
#                     resume_exp = dic["experiance_in_number"]
#                     created_by = dic["created_by"]

#                     # Fetch all JDs from the jd_collection
                    
#                     client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
#                     db = client['resume_parser']
#                     score_resume_by_new_resume = db['score_resume_by_new_resume']
#                     jd_collection = db['jd_collecction']
#                     jds = list(jd_collection.find())

#                     for jd in jds:
#                         try:
#                             JD_skills = jd.get('skills', [])
#                             JD_exp = jd.get('exp', 0.0)
#                             jd_id = jd.get('jd_id', '')

#                             # Calculate score for the resume against the current JD
#                             score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)

#                             scored_resume = {
#                                 'resume_id': str(inserted_id),
#                                 'resume_data': dic,
#                                 'score': score,
#                                 'jd_data': jd,
#                                 'created_by': created_by
#                             }

#                             # Query to find the document
#                             query = {"jd_data.jd_id": jd_id, "resume_id": str(inserted_id)}
#                             # Attempt to find the document
#                             existing_document = score_resume_by_new_jd.find_one(query)

#                             if existing_document:
#                                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
#                                 score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
#                             else:
#                                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
#                                 score_resume_by_new_jd.insert_one(scored_resume)
#                         except Exception as e:
#                             logger.error(f'Error scoring resume against JDs: {str(e)}')
#                     logger.info('Scored resume against all JDs and saved successfully.')
#                 except Exception as e:
#                     logger.error(f'Error scoring resume against JDs: {str(e)}')
#             else:
#                 # Specify the collection name (replace 'mycollection' with your collection name)
#                 file_data = {
#                     "file_name": blob_name,
#                     "req_no": req_no,
#                     "uploaded_by": created_by,
#                     "original_file_name": original_file_name,
#                     "duplicate": "yes"
#                 }
#                 uploaded_files = db["uploaded_files"]
#                 # Insert the JSON data into the collection
#                 uploaded_files.insert_one(file_data)
#             client.close()
#     except Exception as e:
#         print(f'Error :'+str(e))



# def process_job_description(ai_jd_data5):
#     url = 'http://localhost:8005/'
#     try:
                
#         response = requests.get(url)
           
#                 # Check if the request was successful (status code 200)
#         if response.status_code == 200:
#                     # Parse the JSON response
#             data = response.json()
               
#                     # Extracting the 'category' values and storing them in 'Categories'
#             Categories = [item['category'] for item in data]                  
#         else:
#             print(f"Failed to retrieve data. Status code: {response.status_code}")
#     except Exception as e:
#         print(f"An error occurred: {e}")  # Replace with the correct path if needed


def get_matching_resumes_with_scores(json_data):
    try:
        # Extract JD details
        JD_pri_skills = json_data.get('primary_skills', [])  # Expecting an array now
        JD_sec_skills = json_data.get('secondary_skills', [])  # Expecting an array now
        JD_exp = json_data.get('experience_required', 0.0)
        jd_id = json_data.get('jd_id', '')
        jd_title = json_data.get('job_title', '')
        jd_category = json_data.get('category', '')
        created_by = json_data.get('created_by', '')


        # Prepare JD details for processing
        # Prepare JD details for processing
        experience_required = int(JD_exp)
        primary_skills = [skill.strip().lower() for skill in JD_pri_skills] if isinstance(JD_pri_skills, list) else []
        secondary_skills = [skill.strip().lower() for skill in JD_sec_skills] if isinstance(JD_sec_skills, list) else []
        

        


        print(f"JD Details: Experience: {experience_required}, Primary Skills: {primary_skills}, Category: {jd_category}")

        # Connect to MongoDB
        client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
        db = client["resume_parser"]
        # resume_collection = db["resume_collection"]
        # score_resume_by_new_jd = db["score_resume_by_new_jd"]
        print("MongoDB connection successful")
        upload_record = db['upload_record']
        score_resume_by_new_jd = db['score_resume_by_new_jd']
        jd_collection = db['jd_collecction']
        resume_collection = db['resume_collection']
        
        json_data['uploaded_at'] = datetime.datetime.now()
        print("testing2....")
        
        jd_collection.insert_one(json_data)
        
        print("testing1....")
        resumes = list(resume_collection.find())
        print("testing3....")

        # Fetch resumes matching the experience and category
        initial_query = {
            "experiance_in_number": {"$gte": experience_required},
            "Resume_Category": jd_category
        }
        print("Query to fetch resumes:", initial_query)
        potential_resumes = list(resume_collection.find(initial_query))
        print(f"Potential resumes fetched: {len(potential_resumes)}")

        # Filter resumes by skill threshold
        matching_resumes = []
        for resume in potential_resumes:
            resume_skills = {skill.strip().lower() for skill in resume.get("skills", [])}
            matching_skill_count = len(resume_skills.intersection(primary_skills))
            required_skill_threshold = max(1, int(0.1 * len(primary_skills)))  # At least one skill match

            if matching_skill_count >= required_skill_threshold:
                matching_resumes.append(resume)

        print(f"Matching resumes after skill filtering: {len(matching_resumes)}")

        # Calculate scores and save results
        scored_resumes = []
        for resume in matching_resumes:
            score = calculate_score(resume, {
                "experience_required": experience_required,
                "primary_skills": JD_pri_skills,
                "secondary_skills": JD_sec_skills,
                "job_role": json_data.get("job_role", ""),
                "domain_specific": json_data.get("domain_specific", [])
            })

            # scored_resume = {
            #     "jd_data": {
            #         "jd_id": str(jd_id),
            #         "jd_title": jd_title,
            #         "jd_category": jd_category
            #     },
            #     "resume_data": resume,
            #     "created_by": created_by
            # }
            scored_resume = {
                    'resume_id': str(resume['_id']),
                    'resume_data': resume,
                    'score': score,
                    'jd_data': json_data,
                    'created_by': created_by
                }
            scored_resumes.append(scored_resume)

        # Save to MongoDB
        if scored_resumes:
            score_resume_by_new_jd.insert_many(scored_resumes)
            print(f"Scores saved for {len(scored_resumes)} resumes.")
        else:
            print("No scores to save. No matching resumes found.")

    except Exception as e:
        print(f"Error in processing resumes: {e}")



Instruction_pcd = ["""
    You must adhere the below format in the output.
      {
        "project_score": always integer,
        "certifications_score": always integer,
        "domain_specific_score": always integer,
        }
    
    Below iam giving you the sample output format for the reference only.           
    sample Output for reference:
        {
            "project_score": 4,
            "certifications_score": 3,
            "domain_specific_score": 0,
            }
     
"""]
 
agent_pcd = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
              instructions=Instruction_pcd,
              add_references_to_prompt=True)
 
def calculate_score_new(resume, jd_details):
    """Calculate a matching score for a resume based on JD details."""
    try:
        # Extract JD parameters
        experience_required = float(jd_details.get("experience_required", 0))
        primary_skills = {skill.strip().lower() for skill in jd_details.get("primary_skills", "").split(",") if skill}
        secondary_skills = {skill.strip().lower() for skill in jd_details.get("secondary_skills", "").split(",") if skill}
        jd_job_role = jd_details.get("job_role", "").strip().lower()
        primary_skills_count = len(primary_skills)
        secondary_skills_count = len(secondary_skills)
 
        import requests
 
        # Define API endpoint for Score Weights
        score_weights_url = "https://aippoint.ai/aippoint-django-scoringweights/score-weights/"  # Fetch all records
 
        # Fetch data
        response = requests.get(score_weights_url , verify=False)
 
        # Check if the request was successful
        if response.status_code == 200:
            score_data_list = response.json()  # This is a list of dictionaries
           
            if score_data_list:  # Ensure list is not empty
                score_data = score_data_list[0]  # Get the first dictionary
 
                # Store values in variables
                experience_score_weight = score_data["experience_score_weight"]
                project_score_weight = score_data["project_score_weight"]
                certifications_score_weight = score_data["certifications_score_weight"]
                role_score_weight = score_data["role_score_weight"]
 
            else:
                print("No data found in Score Weights API.")
 
        else:
            print("Failed to fetch Score Weights. Status Code:", response.status_code)
 
 
        # Define API endpoint for Skills Score Weights
        skills_score_weights_url = "https://aippoint.ai/aippoint-django-scoringweights/skills-score-weights/"  # Fetch all records
 
        # Fetch data
        response = requests.get(skills_score_weights_url , verify=False)
 
        # Check if the request was successful
        if response.status_code == 200:
            skills_score_data_list = response.json()  # This is a list of dictionaries
           
            if skills_score_data_list:  # Ensure list is not empty
                skills_score_data = skills_score_data_list[0]  # Get the first dictionary
 
                # Store values in variables
                domain_specific_score_weight = skills_score_data["domain_specific_score_weight"]
                skills_score_weight = skills_score_data["skills_score_weight"]
                primary_skills_percentage = skills_score_data["primary_skills_percentage"]
                secondary_skills_percentage = skills_score_data["secondary_skills_percentage"]
 
            else:
                print("No data found in Skills Score Weights API.")
 
        else:
            print("Failed to fetch Skills Score Weights. Status Code:", response.status_code)
 
       
   
        # Extract resume details
        resume_skills = {skill.strip().lower() for skill in resume.get("skills", [])}
        resume_experience = float(resume.get("experiance_in_number", 0))
        resume_job_role = resume.get("job_role", "").strip().lower()
 
        Instruction_ps = """
                        The output format must adhre the following format.
                        {primary_skill_match : always integer
                        secondary_skill_match: always integer}
 
                        the sample output format is below.
                           {"primary_skill_match": 9,
                           "secondary_skill_match": 7}
                                                       """
        agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                        instructions=Instruction_ps,
                        add_references_to_prompt=True)
        run: RunResponse = agent.run(f"""
                                  primary_skill_match : Compare the jd skills {primary_skills} with the resume skills {resume_skills} and give me the number of matching skills.
                                  secondary_skill_match : compare the  jd skills {secondary_skills} with the resume skills {resume_skills} and give me the number of matching skills.
                                    """)
        primary_secondary_match = run.content.strip().strip("```").replace("text", "").strip()
        import re
       # Using regex to extract numbers
        primary_skill_match = int(re.search(r'"primary_skill_match":\s*(\d+)', primary_secondary_match).group(1))
        secondary_skill_match = int(re.search(r'"secondary_skill_match":\s*(\d+)', primary_secondary_match).group(1))
 
        print("primary_skills_match:",primary_skill_match)
        print("secondary_skills_match:",secondary_skill_match)
 
        # Experience Score Calculation
        if experience_required > 0:
            experience_score = min(resume_experience / experience_required, 1.0)
            percentage_difference = (resume_experience - experience_required) / experience_required
 
            # Adjust experience score based on percentage difference
            if percentage_difference >= 0.5:
                experience_score = 0.0
            elif percentage_difference >= 0.4:
                experience_score = max(experience_score - 0.20, 0.0)
            elif percentage_difference >= 0.3:
                experience_score = max(experience_score - 0.15, 0.0)
            elif percentage_difference >= 0.2:
                experience_score = max(experience_score - 0.10, 0.0)
            elif percentage_difference <= -0.4:
                experience_score = max(experience_score - 0.20, 0.0)
            elif percentage_difference <= -0.3:
                experience_score = max(experience_score - 0.15, 0.0)
            elif percentage_difference <= -0.2:
                experience_score = max(experience_score - 0.10, 0.0)
            else:
                experience_score = 1.0  # Full score
        else:
            experience_score = 0.0
        experience_score *= experience_score_weight
 
        # Project and Certification Scoring using the Agent
        project_details = resume.get("projects", [])
        certifications = resume.get("certificates", [])
        domain_specific_resume = resume.get("domain_specific", [])
        domain_specific_jd = jd_details.get("domain_specific", [])
 
        run: RunResponse = agent_pcd.run(f"""give me the below weights based on the provided data.
                                 project_score : 'Compare the {project_details} from the resume with the {jd_details} provided. Analyze how relevant the projects are to the complete 'job description' or 'Primary Skills' and assign a score out of {project_score_weight}. The score should reflect the alignment of the projects with the key requirements or objectives of the job role, such as skills, technologies, or domains mentioned in the JD. Provide the score as an integer between 0 and {project_score_weight}.'
                                 certifications_score : 'Compare the {certifications} from the resume with the {jd_details} provided. Analyze how relevant the certifications are to the 'job description' or 'Primary Skills'and assign a score out of {certifications_score_weight}. The score should reflect the alignment of the certifications with the skills, technologies, or qualifications specified in the JD. Provide the score as an integer between 0 and {certifications_score_weight}.'
                                 domain_specific_score :'Compare the domain-specific relevance of a candidate's {domain_specific_resume} with the Job Description{domain_specific_jd} with a score between 0 and {domain_specific_score_weight}, considering a maximum weightage of {domain_specific_score_weight} points. If domain-specific information is provided in the JD, compare it with the resume and score based on alignment: full oints for an exact match, less for partial relevance, and 0 for no relevance.
                                 If no domain-specific information is mentioned in the JD or 'Not Provided', automatically assign a score of 0'
                                 give me the score outputs in the integer format only.' and overall output in the json format only.
                                   """)
 
        extracted_agent_scores = (run.content).strip().strip("```json").strip()
        # print(extracted_agent_scores)
 
        # Extract scores with error handling
        try:
            # Extract numeric scores
            project_score = int(re.search(r'"project_score":\s*(\d+)', extracted_agent_scores).group(1))
            certifications_score = int(re.search(r'"certifications_score":\s*(\d+)', extracted_agent_scores).group(1))
            domain_specific_score = int(re.search(r'"domain_specific_score":\s*(\d+)', extracted_agent_scores).group(1))
 
            project_rating = min(5, max(1, round((project_score / project_score_weight) * 5))) if project_score_weight > 0 else 1
            certifications_rating = min(5, max(1, round((certifications_score / certifications_score_weight) * 5))) if certifications_score_weight > 0 else 1
            domain_specific_rating = min(5, max(1, round((domain_specific_score / domain_specific_score_weight) * 5))) if domain_specific_score_weight > 0 else 1
 
 
        except (IndexError, ValueError):
            raise ValueError("Failed to parse agent scores. Check the agent output format.")
 
        # Skill Scores
 
 
        primary_skills_score_weight = primary_skills_percentage * skills_score_weight
        secondary_skills_score_weight = secondary_skills_percentage * skills_score_weight
 
        # Skill Scores
        if domain_specific_score == 0:
            primary_skill_score = (primary_skill_match / len(primary_skills)) * primary_skills_score_weight if primary_skills else 0
            secondary_skill_score = (secondary_skill_match / len(secondary_skills)) * secondary_skills_score_weight if secondary_skills else 0
        else:
            primary_skill_score = (primary_skill_match / len(primary_skills)) * 0.8 * primary_skills_score_weight if primary_skills else 0
            secondary_skill_score = (secondary_skill_match / len(secondary_skills)) * secondary_skills_score_weight if secondary_skills else 0
 
 
        # Role Score (Max Weight: 10)
        role_score = role_score_weight if jd_job_role and resume_job_role and jd_job_role == resume_job_role else 0
 
        # Total Score
        total_score = (experience_score + primary_skill_score + secondary_skill_score +
                       role_score + project_score + certifications_score + domain_specific_score)
       
        print(f"Experience Score: {experience_score}")
        print(f"Primary Skills Score: {primary_skill_score}")
        print(f"Secondary Skills Score: {secondary_skill_score}")
        print(f"Role Score: {role_score}")
        print(f"Project Score: {project_score}")
        print(f"Certifications Score: {certifications_score}")
        print(f"Domain-Specific Score: {domain_specific_score}")
 
 
        #Explanations
        if percentage_difference >= 0.5:
            experience_score_expln = "Scored 0 out of 5."
        elif percentage_difference >= 0.4:
            experience_score_expln = f"Scored {max(experience_score - 0.20, 0.0) * 5:.1f} out of 5."
        elif percentage_difference >= 0.3:
            experience_score_expln = f"Scored {max(experience_score - 0.15, 0.0) * 5:.1f} out of 5."
        elif percentage_difference >= 0.2:
            experience_score_expln = f"Scored {max(experience_score - 0.10, 0.0) * 5:.1f} out of 5."
        elif percentage_difference <= -0.4:
            experience_score_expln = f"Scored {max(experience_score - 0.20, 0.0) * 5:.1f} out of 5."
        elif percentage_difference <= -0.3:
            experience_score_expln = f"Scored {max(experience_score - 0.15, 0.0) * 5:.1f} out of 5."
        elif percentage_difference <= -0.2:
            experience_score_expln = f"Scored {max(experience_score - 0.10, 0.0) * 5:.1f} out of 5."
        else:
            experience_score_expln = "Scored 5 out of 5."
 
        # primary_skills_rating = min(5, max(1, round((primary_skill_match / primary_skills_count) * 5))) if primary_skills_count > 0 else 5
        # primary_skills_score_expln = f"Scored {primary_skills_rating} out of 5."
        # # primary_skills_score_expln = f"Score {primary_skills_rating} out of 5 based on {primary_skill_match} out of {primary_skills_count} primary skills matched."
 
        # secondary_skills_rating = min(5, max(1, round((secondary_skill_match / secondary_skills_count) * 5))) if secondary_skills_count > 0 else 5
        # secondary_skills_score_expln = f"Scored {secondary_skills_rating} out of 5."
        # # secondary_skills_score_expln = f"Score {secondary_skills_rating} out of 5 based on {secondary_skill_match} out of {secondary_skills_count} secondary skills matched."
       
        # domain_specific_score_explanation = f"Scored {domain_specific_rating} out of 5."
        # # domain_specific_score_explanation = f"Score {domain_specific_rating} out of 5 based on domain-specific knowledge evaluation with a weight of {domain_specific_score_weight}."
        # Primary Skills Score Explanation
        if domain_specific_score == 0:
            primary_skills_rating = (
                min(5, max(1, round((primary_skill_match / len(primary_skills)) * 5)))
                if primary_skills else 1
            )
            secondary_skills_rating = (
                min(5, max(1, round((secondary_skill_match / len(secondary_skills)) * 5)))
                if secondary_skills else 1
            )
        else:
            primary_skills_rating = (
                min(5, max(1, round((primary_skill_match / len(primary_skills)) * 5)))
                if primary_skills else 1
            )
            secondary_skills_rating = (
                min(5, max(1, round((secondary_skill_match / len(secondary_skills)) * 5)))
                if secondary_skills else 1
            )
 
        # Explanations
        primary_skills_score_expln = f"Scored {primary_skills_rating} out of 5."
        secondary_skills_score_expln = f"Scored {secondary_skills_rating} out of 5."
 
        # Domain-Specific Skills Score Explanation (Only if domain-specific score is present)
        if domain_specific_score > 0:
            domain_specific_rating = min(5, max(1, round((domain_specific_score / domain_specific_score_weight) * 5)))
            domain_specific_score_explanation = f"Scored {domain_specific_rating} out of 5."
        else:
            domain_specific_score_explanation = "Not applicable."
 
 
        role_rating = 5 if resume_job_role == jd_job_role else 1
        role_score_expln = f"Scored {role_rating} out of 5."
        # role_score_expln = f"Score {role_rating} out of 5 as the job role {'matches' if role_rating == 5 else 'does not match'} the JD."
 
        project_score_explanation = f"Scored {project_rating} out of 5."
        # project_score_explanation = f"Scored {project_rating} out of 5 based on project evaluation with a weight of {project_score_weight}."
       
        certifications_score_explanation = f"Scored {certifications_rating} out of 5."
        # certifications_score_explanation = f"Score {certifications_rating} out of 5 based on certification evaluation with a weight of {certifications_score_weight}."
       
 
 
 
        resume_score_explanations = {
        "experience": experience_score_expln,
        "primary_skills": primary_skills_score_expln,
        "secondary_skills": secondary_skills_score_expln,
        "role": role_score_expln,
        "projects": project_score_explanation,
        "certifications": certifications_score_explanation,
        "domain_specific": domain_specific_score_explanation
    }
 
        print("Explanations for the Scores:")
        for key, explanation in resume_score_explanations.items():
            print(f"{key.capitalize()} Score Explanation: {explanation}")
 
 
        # # Debugging Output
        # print(f"Scores:\n Experience: {experience_score:.2f}, Primary Skill: {primary_skill_score:.2f}, "
        #       f"Secondary Skill: {secondary_skill_score:.2f}, Role: {role_score:.2f}, "
        #       f"Project: {project_score:.2f}, Certifications: {certifications_score:.2f}, "
        #       f"Domain-Specific: {domain_specific_score:.2f}")
 
        return total_score ,resume_score_explanations
    except Exception as e:
        print(f"Error calculating score: {e}")
        return 0


def print_file_paths_and_convert_docx(data_obj):
    try:
        blob_name = data_obj["file_name"]
        created_by = data_obj["created_by"]
        req_no = data_obj["req_no"]
        original_file_name = data_obj["original_file_name"]
        if blob_name.endswith(".pdf"):
            
            data = pdf_to_text(blob_name)
            pdf_page_content = data.get("pdf_page_content", []) 
            # Ensure the extracted page content is not empty
            if not pdf_page_content:
                print("No content extracted from the PDF.")
                return

            # Create Document objects from page content for embedding
            documents = [LangchainDocument(page_content=text) for text in pdf_page_content]

            url = 'https://aippoint.ai/aippoint-spring-category/'  # Replace with the correct path if needed
           
            try:
                # Making a GET request to the API
                response = requests.get(url, verify=False)
           
                # Check if the request was successful (status code 200)
                if response.status_code == 200:
                    # Parse the JSON response
                    data = response.json()
               
                    # Extracting the 'category' values and storing them in 'Categories'
                    Categories = [item['category'] for item in data]
               
                   
                else:
                    print(f"Failed to retrieve data. Status code: {response.status_code}")
            except Exception as e:
                print(f"An error occurred: {e}")

            
            try:
                # instructor_embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
                embedding = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
                print("GoogleGenerativeAI embeddings initialized successfully.")
            except Exception as e:
                print(f"Error initializing GoogleGenerativeAI embeddings: {e}")
                # Optionally, log the traceback for deeper debugging insights
                import traceback
                print(traceback.format_exc())
                return

            try:
                # Embed the documents using FAISS and instructor embeddings
                # vectordb = FAISS.from_documents(documents=documents, embedding=instructor_embeddings)
                # retriever = vectordb.as_retriever(score_threshold=0.7)
                # Embed each chunk and load it into the FAISS vector store
                faiss_index_path = "faiss_index"
                db = FAISS.from_documents(documents, embedding)

                # Save the FAISS index to disk
                db.save_local(faiss_index_path)

                # Load the FAISS index from disk
                db = FAISS.load_local(faiss_index_path, embedding, allow_dangerous_deserialization=True) 

                # Create a retriever from the vector store
                retriever = db.as_retriever()

                # Create a knowledge base from the vector store
                knowledge_base = LangChainKnowledgeBase(retriever=retriever)
                print("FAISS database and retriever setup completed.")

            except Exception as e:
                print(f"Error setting up FAISS database: {e}")
                return
           

            Instruction =[ """
                            The output should consistently adhere to the following format provided below:
                            {
                                "name": "",
                                "mobile_number": "",
                                "email_id": "",
                                "linkedin_id": "",
                                "location": "",
                                "skills": "",
                                "certifications": "",
                                "experience_in_number": this value is always integer,
                                "job_role" :"",
                                "domain_specific" :"",
                                "experience_details": [
                                    {
                                        "designation_at_company": "",
                                        "company": "",
                                        "duration": ""
                                    },
                                    {
                                        "designation_at_company": "",
                                        "company": "",
                                        "duration": ""
                                    }
                                ],
                                "education_details": [
                                    {
                                        "Degree": "",
                                        "institution": "",
                                        "year of graduation": ""
                                    }
                                ],
                                "project_details": [
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    },
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    },
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    }
                                ],
                                "resume_category": ""
                            }

                            """]

            agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                          instructions = Instruction,
                          knowledge_base=knowledge_base, add_references_to_prompt=True)
            

            run: RunResponse = agent.run(f"""Extract the following details from the {pdf_page_content} uploded resume. Ensure they match the instruction format:
                             
                                        name: "Extract the full name of the individual from the provided details."
                                        mobile_number: "Extract the mobile number provided, including the country code and other details."
                                        email_id: "Identify and extract the email address from the provided details."
                                        linkedin_id: "Extract the LinkedIn profile URL associated with the individual from the given details."
                                        location: "Determine the location (city, country, or region) of the individual based on the information provided."
                                        experience_in_number: "Extract the number of years of experience the individual has, as specified. This value should always be in integer format."
                                        skills: "Identify and list all the skills the individual possesses as mentioned in the provided details. These skills may cover various areas such as technical, soft skills, or domain expertise."
                                        certifications: "Check if any certifications are mentioned for the individual. If provided, extract the certification details; otherwise, mark as 'Not Provided'."
                                        job_role: "Extract the most recent job title or designation held by the individual, as mentioned in the 'work' section of their details. If multiple roles are listed, prioritize based on the most recent start date or the designation marked as 'current.' 
                                                    If the job title is not explicitly mentioned, analyze the resume comprehensively, including project descriptions, skills, and experience details, to infer the most likely job role based on the information provided."
                                        domain_specific :"Extract the domain-specific experience mentioned in the resume. Focus on identifying industries (e.g., healthcare, e-commerce, finance), platforms, or products the candidate has worked on, as explicitly stated in the resume or in the 'projects' section.
                                                          If multiple domains are present, include them in a list. Ensure to analyze the 'projects' section for any additional domain-specific details. Provide the extracted domain-specific details". 
                                        experience_details: in the output this is always the list of arrays.
                                        designation_at_company: "Extract the specific job role or designation the individual held at each company mentioned in their experience."
                                        company: "Identify and extract the name of the company the individual worked for during their experience."
                                        duration: "Extract the start and end dates (or current status if ongoing) of the individual’s tenure at each job."
                                        education_details:
                                        Degree: "Identify and extract the degree(s) obtained by the individual, such as Bachelor's, Master's, etc."
                                        institution: "Extract the name of the educational institution where the individual completed their degree."
                                        year of graduation: "Extract the year when the individual graduated from the educational institution."
                                        project_details:
                                        project_name: "Extract the name or title of the project the individual worked on."
                                        description: "Provide a brief overview of the project, describing its purpose, scope, and impact."
                                        roles_and_responsibilities: "Extract the roles and responsibilities the individual held within the project."
                                        tools_and_technologies: "List the tools, technologies, or programming languages used in the project by the individual."
                                        team_size: "Extract the team size or mention 'Not Provided' if team size is not specified."
                                        resume_category -"Perform a comprehensive analysis of the provided resume, taking into account all relevant information such as skills, projects, 
                                         and the job role. Based on this analysis, strictly match the resume with the most appropriate category from the predefined {Categories}. 
                                         If no suitable category is found in the predefined list, assign the category as 'Not Provided'."
                                        Do not include any explanations or introductory lines. The output format contains only the category. Print everything in regular text (not bold or italic)"
                            give me the output in the string format.""")
            resume_info = run.content

            import re
            # Clean up the string (remove triple backticks and "text")
            extracted_details = resume_info.strip().strip("```json").strip()
            print(extracted_details)

            # Extract name
            name_match = re.search(r'"name":\s*"([^"]+)"', extracted_details)
            name = name_match.group(1) if name_match else "Name not found"

            # Extract mobile number
            mobile_number_match = re.search(r'"mobile_number":\s*"([^"]+)"', extracted_details)
            mobile_number = mobile_number_match.group(1) if mobile_number_match else "Mobile number not found"

            # Extract LinkedIn ID
            linkedin_id_match = re.search(r'"linkedin_id":\s*"([^"]+)"', extracted_details)
            linkedin_id = linkedin_id_match.group(1) if linkedin_id_match else "LinkedIn ID not found"

            # Extract email
            email_id_match = re.search(r'"email_id":\s*"([^"]+)"', extracted_details)
            email_id = email_id_match.group(1) if email_id_match else "Email ID not found"

            # Extract location
            location_match = re.search(r'"location":\s*"([^"]+)"', extracted_details)
            location = location_match.group(1) if location_match else "Location not found"

            # Extract job_role
            job_role_match = re.search(r'"job_role":\s*"([^"]+)"', extracted_details)
            job_role = job_role_match.group(1) if job_role_match else "Location not found"

            domain_specific_match = re.search(r'"domain_specific":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            domain_specific = [domain.strip().strip('"') for domain in domain_specific_match.group(1).split(",")] if domain_specific_match else ["Domain Specific not found"]

            # Extract skills
            skills_match = re.search(r'"skills":\s*"([^"]+)"', extracted_details)
            skills = skills_match.group(1).split(",") if skills_match else ["Skills not found"]

            # Extract certifications
            certifications_match = re.search(r'"certifications":\s*"([^"]+)"', extracted_details)
            certifications = certifications_match.group(1) if certifications_match else "Certifications not found"
            
            # experience_match = re.search(r'"experience_in_number":\s*"([^"]+)"', extracted_details)
            # experience_in_number = experience_match.group(1) if experience_match else "Experience not found"

            # Correct regex to match a raw number
            experience_match = re.search(r'"experience_in_number":\s*(\d+)', extracted_details)
            experience_in_num = experience_match.group(1) if experience_match else "Experience not found"
            experience_in_number = int(experience_in_num)


            # Extract education details (corrected regex for arrays)
            education_match = re.search(r'"education_details":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            education = education_match.group(1).strip() if education_match else "Education not found"

           

            work_match = re.search(r'"experience_details":\s*(\[[^\]]*\])', extracted_details, re.DOTALL)

            # Parse the matched array if found
            if work_match:
                work_array = json.loads(work_match.group(1))  # Convert JSON string to a Python list
            else:
                work_array = "Work not found"
            print("array",work_array)
            work=work_array

            # Extract project details (corrected regex for arrays)
            project_details_match = re.search(r'"project_details":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            project_details = project_details_match.group(1).strip() if project_details_match else "Project Details not found"

            # Extract resume category
            resume_category_match = re.search(r'"resume_category":\s*"([^"]+)"', extracted_details)
            resume_category = resume_category_match.group(1) if resume_category_match else "Resume Category not found"


            dic = {
                "name": name,
                "email": email_id,
                "phone": mobile_number,
                "experiance_in_number":experience_in_number,
                "url": linkedin_id,
                "location": location,
                "job_role": job_role,
                "domain_specific":domain_specific,
                "work": work,
                "education": education,
                "certificates": certifications,
                "skills": skills, 
                "projects": project_details,
                "created_by": created_by,
                "req_no": req_no,
                "Resume_Category":resume_category
            } 

            dic['file_name'] = blob_name
            client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
            db = client["resume_parser"]
            collection = db["resume_collection"]
            existing_document = collection.find_one({"email": {"$regex": re.compile('^' + re.escape(dic['email']) + '$', re.IGNORECASE)}})
            if existing_document is None:
                # If the document doesn't exist, insert the new data into the collection
                collection.insert_one(dic)
                inserted_id = collection.inserted_id
                file_data = {
                    "file_name": blob_name,
                    "req_no": req_no,
                    "uploaded_by": created_by,
                    "original_file_name": original_file_name,
                    "duplicate": "no"
                }
                uploaded_files = db["uploaded_files"]
                # Insert the JSON data into the collection
                uploaded_files.insert_one(file_data)
                try:
                    # Process resume data
                    resume_skills = dic["skills"]
                    resume_exp = dic["experiance_in_number"]
                    created_by = dic["created_by"]
                   
                    client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
                    db = client['resume_parser']
                    score_resume_by_new_jd = db['score_resume_by_new_jd']
                    jd_collection = db['jd_collecction']
 
                    jds = list(jd_collection.find())
 
                    for jd in jds:
                        try:
                            JD_skills = jd.get('skills', [])
                            JD_exp = jd.get('exp', 0.0)
                            jd_id = jd.get('jd_id', '')
 
                            # Calculate score for the resume against the current JD
                            # score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
 
                            scored_resume = {
                                'resume_id': str(inserted_id),
                                'resume_data': dic,
                                'score': score,
                                'jd_data': jd,
                                'created_by': created_by
                            }
 
                            query = {"jd_data.jd_id": jd_id, "resume_id": str(inserted_id)}
                            # Attempt to find the document
                            existing_document = score_resume_by_new_jd.find_one(query)
 
                            if existing_document:
                                # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
                                score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
                            else:
                                # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
                                score_resume_by_new_jd.insert_one(scored_resume)
                        except Exception as e:
                            logger.error(f'Error scoring resume against JDs: {str(e)}')
                    logger.info('Scored resume against all JDs and saved successfully.')
                except Exception as e:
                    logger.error(f'Error scoring resume against JDs: {str(e)}')
            else:
                # Specify the collection name (replace 'mycollection' with your collection name)
                file_data = {
                    "file_name": blob_name,
                    "req_no": req_no,
                    "uploaded_by": created_by,
                    "original_file_name": original_file_name,
                    "duplicate": "yes"
                }
                uploaded_files = db["uploaded_files"]
                # Insert the JSON data into the collection
                uploaded_files.insert_one(file_data)
            client.close()
 
 
        elif blob_name.lower().endswith((".doc", ".docx", ".dox")):
            doc_data = docx_to_text(blob_name)
            
            if doc_data is None:
                print("Error extracting content from the docx file.")
                return
 
            doc_page_content = doc_data.get("doc_page_content", [])
 
            # Ensure the extracted page content is not empty
            if not doc_page_content:
                print("No content extracted from the docx.")
                return
            #here
            print("doc_page_content",doc_page_content)

            documents = [LangchainDocument(page_content=text, metadata={"source": "doc_page_content"}) for text in doc_page_content]
            print("DOCSSS",documents)
            url = 'https://aippoint.ai/aippoint-spring-category/'  # Replace with the correct path if needed
           
            try:
                # Making a GET request to the API
                response = requests.get(url, verify=False)
           
                # Check if the request was successful (status code 200)
                if response.status_code == 200:
                    # Parse the JSON response
                    data = response.json()
               
                    # Extracting the 'category' values and storing them in 'Categories'
                    Categories = [item['category'] for item in data]
               
                   
                else:
                    print(f"Failed to retrieve data. Status code: {response.status_code}")
            except Exception as e:
                print(f"An error occurred: {e}")

            
            try:
                # instructor_embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
                embedding = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
                print("GoogleGenerativeAI embeddings initialized successfully.")
            except Exception as e:
                print(f"Error initializing GoogleGenerativeAI embeddings: {e}")
                # Optionally, log the traceback for deeper debugging insights
                import traceback
                print(traceback.format_exc())
                return

            try:
                # Embed the documents using FAISS and instructor embeddings
                # vectordb = FAISS.from_documents(documents=documents, embedding=instructor_embeddings)
                # retriever = vectordb.as_retriever(score_threshold=0.7)
                # Embed each chunk and load it into the FAISS vector store
                faiss_index_path = "faiss_index"
                db = FAISS.from_documents(documents, embedding)

                # Save the FAISS index to disk
                db.save_local(faiss_index_path)

                # Load the FAISS index from disk
                db = FAISS.load_local(faiss_index_path, embedding, allow_dangerous_deserialization=True) 

                # Create a retriever from the vector store
                retriever = db.as_retriever()

                # Create a knowledge base from the vector store
                knowledge_base = LangChainKnowledgeBase(retriever=retriever)
                print("FAISS database and retriever setup completed.")

            except Exception as e:
                print(f"Error setting up FAISS database: {e}")
                return
           

            Instruction =[ """
                            The output should consistently adhere to the following format provided below:
                            {
                                "name": "",
                                "mobile_number": "",
                                "email_id": "",
                                "linkedin_id": "",
                                "location": "",
                                "skills": "",
                                "certifications": "",
                                "experience_in_number": this value is always integer,
                                "job_role" :"",
                                "domain_specific" :"",
                                "experience_details": [
                                    {
                                        "designation_at_company": "",
                                        "company": "",
                                        "duration": ""
                                    },
                                    {
                                        "designation_at_company": "",
                                        "company": "",
                                        "duration": ""
                                    }
                                ],
                                "education_details": [
                                    {
                                        "Degree": "",
                                        "institution": "",
                                        "year of graduation": ""
                                    }
                                ],
                                "project_details": [
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    },
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    },
                                    {
                                        "project_name": "",
                                        "description": "",
                                        "roles_and_responsibilities": "",
                                        "tools_and_technologies": "",
                                        "team_size": ""
                                    }
                                ],
                                "resume_category": ""
                            }

                            """]

            agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                          instructions = Instruction,
                          knowledge_base=knowledge_base, add_references_to_prompt=True)
            

            run: RunResponse = agent.run(f"""Extract the following details from the {doc_page_content} uploded resume. Ensure they match the instruction format:
                             
                                        name: "Extract the full name of the individual from the provided details."
                                        mobile_number: "Extract the mobile number provided, including the country code and other details."
                                        email_id: "Identify and extract the email address from the provided details."
                                        linkedin_id: "Extract the LinkedIn profile URL associated with the individual from the given details."
                                        location: "Determine the location (city, country, or region) of the individual based on the information provided."
                                        experience_in_number: "Extract the number of years of experience the individual has, as specified. This value should always be in integer format."
                                        skills: "Identify and list all the skills the individual possesses as mentioned in the provided details. These skills may cover various areas such as technical, soft skills, or domain expertise."
                                        certifications: "Check if any certifications are mentioned for the individual. If provided, extract the certification details; otherwise, mark as 'Not Provided'."
                                        job_role: "Extract the most recent job title or designation held by the individual, as mentioned in the 'work' section of their details. If multiple roles are listed, prioritize based on the most recent start date or the designation marked as 'current.' 
                                                    If the job title is not explicitly mentioned, analyze the resume comprehensively, including project descriptions, skills, and experience details, to infer the most likely job role based on the information provided."
                                        domain_specific :"Extract the domain-specific experience mentioned in the resume. Focus on identifying industries (e.g., healthcare, e-commerce, finance), platforms, or products the candidate has worked on, as explicitly stated in the resume or in the 'projects' section.
                                                          If multiple domains are present, include them in a list. Ensure to analyze the 'projects' section for any additional domain-specific details. Provide the extracted domain-specific details". 
                                        experience_details: in the output this is always the list of arrays.
                                        designation_at_company: "Extract the specific job role or designation the individual held at each company mentioned in their experience."
                                        company: "Identify and extract the name of the company the individual worked for during their experience."
                                        duration: "Extract the start and end dates (or current status if ongoing) of the individual’s tenure at each job."
                                        education_details:
                                        Degree: "Identify and extract the degree(s) obtained by the individual, such as Bachelor's, Master's, etc."
                                        institution: "Extract the name of the educational institution where the individual completed their degree."
                                        year of graduation: "Extract the year when the individual graduated from the educational institution."
                                        project_details:
                                        project_name: "Extract the name or title of the project the individual worked on."
                                        description: "Provide a brief overview of the project, describing its purpose, scope, and impact."
                                        roles_and_responsibilities: "Extract the roles and responsibilities the individual held within the project."
                                        tools_and_technologies: "List the tools, technologies, or programming languages used in the project by the individual."
                                        team_size: "Extract the team size or mention 'Not Provided' if team size is not specified."
                                        resume_category - "From the provided resume data, extract the job title or job role that the resume is associated with,
                                            analyze the entire resume, including the skills and projects, and categorize the resume based on this information
                                            and match with the {Categories}.
                                            Do not include any explanations or introductory lines. The output format contains only the category. Print everything in regular text (not bold or italic)"
                            give me the output in the string format.""")
            resume_info = run.content

            import re
            # Clean up the string (remove triple backticks and "text")
            extracted_details = resume_info.strip().strip("```json").strip()
            print("extracted_details",extracted_details)

            # Extract name
            name_match = re.search(r'"name":\s*"([^"]+)"', extracted_details)
            name = name_match.group(1) if name_match else "Name not found"

            # Extract mobile number
            mobile_number_match = re.search(r'"mobile_number":\s*"([^"]+)"', extracted_details)
            mobile_number = mobile_number_match.group(1) if mobile_number_match else "Mobile number not found"

            # Extract LinkedIn ID
            linkedin_id_match = re.search(r'"linkedin_id":\s*"([^"]+)"', extracted_details)
            linkedin_id = linkedin_id_match.group(1) if linkedin_id_match else "LinkedIn ID not found"

            # Extract email
            email_id_match = re.search(r'"email_id":\s*"([^"]+)"', extracted_details)
            email_id = email_id_match.group(1) if email_id_match else "Email ID not found"

            # Extract location
            location_match = re.search(r'"location":\s*"([^"]+)"', extracted_details)
            location = location_match.group(1) if location_match else "Location not found"

            # Extract job_role
            job_role_match = re.search(r'"job_role":\s*"([^"]+)"', extracted_details)
            job_role = job_role_match.group(1) if job_role_match else "Location not found"

            domain_specific_match = re.search(r'"domain_specific":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            domain_specific = [domain.strip().strip('"') for domain in domain_specific_match.group(1).split(",")] if domain_specific_match else ["Domain Specific not found"]

            # Extract skills
            skills_match = re.search(r'"skills":\s*"([^"]+)"', extracted_details)
            skills = skills_match.group(1).split(",") if skills_match else ["Skills not found"]

            # Extract certifications
            certifications_match = re.search(r'"certifications":\s*"([^"]+)"', extracted_details)
            certifications = certifications_match.group(1) if certifications_match else "Certifications not found"
            
            # experience_match = re.search(r'"experience_in_number":\s*"([^"]+)"', extracted_details)
            # experience_in_number = experience_match.group(1) if experience_match else "Experience not found"

            # Correct regex to match a raw number
            experience_match = re.search(r'"experience_in_number":\s*(\d+)', extracted_details)
            experience_in_num = experience_match.group(1) if experience_match else "Experience not found"
            experience_in_number = int(experience_in_num)


            # Extract education details (corrected regex for arrays)
            education_match = re.search(r'"education_details":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            education = education_match.group(1).strip() if education_match else "Education not found"

           

            work_match = re.search(r'"experience_details":\s*(\[[^\]]*\])', extracted_details, re.DOTALL)

            # Parse the matched array if found
            if work_match:
                work_array = json.loads(work_match.group(1))  # Convert JSON string to a Python list
            else:
                work_array = "Work not found"
            print("array",work_array)
            work=work_array

            # Extract project details (corrected regex for arrays)
            project_details_match = re.search(r'"project_details":\s*\[(.*?)\]', extracted_details, re.DOTALL)
            project_details = project_details_match.group(1).strip() if project_details_match else "Project Details not found"

            # Extract resume category
            resume_category_match = re.search(r'"resume_category":\s*"([^"]+)"', extracted_details)
            resume_category = resume_category_match.group(1) if resume_category_match else "Resume Category not found"


            dic = {
                "name": name,
                "email": email_id,
                "phone": mobile_number,
                "experiance_in_number":experience_in_number,
                "url": linkedin_id,
                "location": location,
                "job_role": job_role,
                "domain_specific":domain_specific,
                "work": work,
                "education": education,
                "certificates": certifications,
                "skills": skills, 
                "projects": project_details,
                "created_by": created_by,
                "req_no": req_no,
                "Resume_Category":resume_category
            } 

            dic['file_name'] = blob_name
            client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
            db = client["resume_parser"]
            collection = db["resume_collection"]
            existing_document = collection.find_one({"email": {"$regex": re.compile('^' + re.escape(dic['email']) + '$', re.IGNORECASE)}})
            if existing_document is None:
                # If the document doesn't exist, insert the new data into the collection
                collection.insert_one(dic)
                inserted_id = collection.inserted_id
                file_data = {
                    "file_name": blob_name,
                    "req_no": req_no,
                    "uploaded_by": created_by,
                    "original_file_name": original_file_name,
                    "duplicate": "no"
                }
                uploaded_files = db["uploaded_files"]
                # Insert the JSON data into the collection
                uploaded_files.insert_one(file_data)
                try:
                    # Process resume data
                    resume_skills = dic["skills"]
                    resume_exp = dic["experiance_in_number"]
                    created_by = dic["created_by"]
                   
                    client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
                    db = client['resume_parser']
                    score_resume_by_new_jd = db['score_resume_by_new_jd']
                    jd_collection = db['jd_collecction']
 
                    jds = list(jd_collection.find())
 
                    for jd in jds:
                        try:
                            JD_skills = jd.get('skills', [])
                            JD_exp = jd.get('exp', 0.0)
                            jd_id = jd.get('jd_id', '')
 
                            # Calculate score for the resume against the current JD
                            # score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
 
                            scored_resume = {
                                'resume_id': str(inserted_id),
                                'resume_data': dic,
                                'score': score,
                                'jd_data': jd,
                                'created_by': created_by
                            }
 
                            query = {"jd_data.jd_id": jd_id, "resume_id": str(inserted_id)}
                            # Attempt to find the document
                            existing_document = score_resume_by_new_jd.find_one(query)
 
                            if existing_document:
                                # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
                                score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
                            else:
                                # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
                                score_resume_by_new_jd.insert_one(scored_resume)
                        except Exception as e:
                            logger.error(f'Error scoring resume against JDs: {str(e)}')
                    logger.info('Scored resume against all JDs and saved successfully.')
                except Exception as e:
                    logger.error(f'Error scoring resume against JDs: {str(e)}')
            else:
                # Specify the collection name (replace 'mycollection' with your collection name)
                file_data = {
                    "file_name": blob_name,
                    "req_no": req_no,
                    "uploaded_by": created_by,
                    "original_file_name": original_file_name,
                    "duplicate": "yes"
                }
                uploaded_files = db["uploaded_files"]
                # Insert the JSON data into the collection
                uploaded_files.insert_one(file_data)
            client.close()
        else:
            print(f'Resume or JD data is empty')
    except Exception as e:
        print(f'Error in processing data here:' + str(e))


 
# def update_answer(reference_number, question_text, answer):
#     print("inputdata", reference_number, question_text, answer)
#     client = MongoClient(f"mongodb://localhost:27017/")
#     db = client["resume_parser"]
#     collection = db["interview_collection"]

#     # Find the document with the matching reference number
#     document = collection.find_one({"reference_number": reference_number})

#     if document:
#         # Update the answer for the specified question text
#         for question in document["questions"]:
#             if question.get("question") == question_text:
#                 question["answer"] = answer
#                 break
#         else:
#             # If no matching question is found
#             return False

#         # Update the entire document in the collection
#         collection.update_one({"_id": document["_id"]}, {"$set": {"questions": document["questions"]}})
#         return True  # Indicate successful update
#     else:
#         return False  # Indicate document not found



def update_answer(object_id_str, question_text, answer):
    
    # Convert the string object ID to an actual ObjectId instance
    object_id = ObjectId(object_id_str)

    client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
 
    db = client["resume_parser"]
    collection = db["interview_collection"]

    # Find the document with the matching ObjectId
    document = collection.find_one({"_id": object_id})
    

    if document:
        for question in document["questions"]:
            if question.get("question") == question_text:
                question["answer"] = answer
                break
        else:
            return False

        # Update the entire document in the collection
        collection.update_one({"_id": document["_id"]}, {"$set": {"questions": document["questions"]}})
        return True  
    else:
        return False 


@shared_task
def celery_upload_file(uploaded_files):
    try:
        print("came here")
        num_threads = 20
        with concurrent.futures.ThreadPoolExecutor(max_workers=num_threads) as executor:
            parsed_results = list(executor.map(print_file_paths_and_convert_docx, uploaded_files))
        executor.shutdown()
        channel_layer = get_channel_layer()
        req_no = uploaded_files[0]["req_no"]
        print("This is the reqno",req_no)
        created_by = uploaded_files[0]["created_by"]
        print("created",created_by)
        sanitized_created_by = re.sub(r'[^a-zA-Z0-9._-]', '_', created_by)[:99]

        # Formulate the group name
        group_name = f'user_notifications_{created_by}'

        print(group_name)  # Check the result
        # # group_name = f'user_notifications_{created_by}'
        async_to_sync(channel_layer.group_send)(
            group_name,
            {
                "type": "user_notification",
                "req_no": req_no,
                "message": f"Your resumes has been processed",
            }
        )
    except Exception as e:
        print(f'Error in celery processing here:'+str(e))


@shared_task
def celery_upload_interview_data(uploaded_files):
    try:
        num_threads = 20
        with concurrent.futures.ThreadPoolExecutor(max_workers=num_threads) as executor:
            parsed_results = list(executor.map(file_paths_and_convert_docx, uploaded_files))
        executor.shutdown()
        channel_layer = get_channel_layer()
        created_by = uploaded_files[0]["created_by"]
        group_name = f'user_notifications_{created_by}'
        async_to_sync(channel_layer.group_send)(
            group_name,
            {
                "type": "user_notification",
                "message": f"Your resumes have been processed",
            }
        )

    except Exception as e:
        print(f'Error in celery processing here:' + str(e))


from celery import shared_task
from pymongo import MongoClient
import json
import google.generativeai as palm
import logging
logger = logging.getLogger(__name__)
 

from celery import shared_task
import json
import time
import logging

logger = logging.getLogger(__name__)

# @shared_task
# def score_resumes_task_by_new_jd(json_data):
  
   
    # try:
    #     JD_skills = json_data.get('skills', [])
    #     JD_exp = json_data.get('exp', 0.0)
    #     jd_id = json_data.get('jd_id', 0)
    #     created_by = json_data.get('created_by', '')
    #     client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
    #     db = client['resume_parser']
    #     upload_record = db['upload_record']
    #     score_resume_by_new_jd = db['score_resume_by_new_jd']
    #     jd_collection = db['jd_collecction']
    #     resume_collection = db['resume_collection']
    #     json_data['uploaded_at'] = datetime.datetime.now()
    #     print("testing2....")
    #     # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
    #     jd_collection.insert_one(json_data)
    #     # Assuming you have a MongoDB collection named 'resume_collection'
    #     # resumes = list(upload_record.find())
    #     print("testing1....")
    #     resumes = list(resume_collection.find())
    #     print("testing3....")
    #     print(resumes.count)
    #     for resume in resumes:
    #         print("testing repeat...")
    #         try:
    #             resume_skills = resume.get('skills', [])
    #             print(resume_skills,"resume skill test...")
    #             resume_exp = resume.get('experiance_in_number', 0.0)

    #             score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
    #             print("calculated score",score)

    #             scored_resume = {
    #                 'resume_id': str(resume['_id']),
    #                 'resume_data': resume,
    #                 'score': score,
    #                 'jd_data': json_data,
    #                 'created_by': created_by
    #             }
    #             # Query to find the document
    #             query = {"jd_data.jd_id": jd_id, "resume_id": str(resume['_id'])}
    #             # Attempt to find the document
    #             existing_document = score_resume_by_new_jd.find_one(query)
    #             print("document",existing_document)
    #             if existing_document:
    #                 print("updateddddddddddddddddd")
    #                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
    #                 score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
    #             else:
    #                 # Assuming you have a MongoDB collection named 'score_resume_by_new_jd'
    #                 score_resume_by_new_jd.insert_one(scored_resume)
    #         except Exception as e:
    #             logger.error(f'Error scoring resume against JDs: {str(e)}')

    #     logger.info('Scored resumes and saved successfully.')
    # except Exception as e:
    #     logger.error(f'Error scoring resumes: {str(e)}')
    #  try:
    #     JD_skills = json_data.get('skills', [])
    #     JD_exp = json_data.get('exp', 0.0)
    #     jd_id = json_data.get('jd_id', 0)
    #     created_by = json_data.get('created_by', '')
    #     JD_category = json_data.get('category', '')

    #     client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
 
    #     db = client['resume_parser']
    #     upload_record = db['upload_record']
    #     score_resume_by_new_jd = db['score_resume_by_new_jd']
    #     jd_collection = db['jd_collecction']
    #     resume_collection = db['resume_collection']
        
    #     json_data['uploaded_at'] = datetime.datetime.now()
    #     print("testing2....")
        
    #     jd_collection.insert_one(json_data)
        
    #     print("testing1....")
    #     resumes = list(resume_collection.find())
    #     print("testing3....")
        
    #     for resume in resumes:
    #         print("testing repeat...")
    #         try:
    #             # Filter resumes by category to match the JD's category
    #             resume_category = resume.get('category', '')
    #             if resume_category != JD_category:
    #                 print(f"Skipping resume with category '{resume_category}', doesn't match JD category '{JD_category}'")
    #                 continue  # Skip this resume if the category doesn't match
                
    #             resume_skills = resume.get('skills', [])
    #             print(resume_skills, "resume skill test...")
    #             resume_exp = resume.get('experiance_in_number', 0.0)

    #             score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
    #             print("calculated score", score)

    #             scored_resume = {
    #                 'resume_id': str(resume['_id']),
    #                 'resume_data': resume,
    #                 'score': score,
    #                 'jd_data': json_data,
    #                 'created_by': created_by
    #             }
                
    #             # Query to find the document
    #             query = {"jd_data.jd_id": jd_id, "resume_id": str(resume['_id'])}
                
    #             # Attempt to find the document
    #             existing_document = score_resume_by_new_jd.find_one(query)
    #             print("document", existing_document)
    #             if existing_document:
    #                 print("updateddddddddddddddddd")
    #                 score_resume_by_new_jd.update_one(query, {"$set": scored_resume})
    #             else:
    #                 score_resume_by_new_jd.insert_one(scored_resume)
    #         except Exception as e:
    #             logger.error(f'Error scoring resume against JDs: {str(e)}')

    #     logger.info('Scored resumes and saved successfully.')
    #  except Exception as e:
    #     logger.error(f'Error scoring resumes: {str(e)}')

import google.generativeai as genai
genai.configure(api_key=GOOGLE_API_KEY)
# Set up the model
generation_config = {
  "temperature": 1,
  "top_p": 0.95,
  "top_k": 40,
  "max_output_tokens": 8192,
  "response_mime_type": "text/plain",
}
model = genai.GenerativeModel(model_name="gemini-1.5-flash",
generation_config=generation_config,)


# Improved calculation with more weight on experience
# def calculate_score(resume_skills, resume_exp, JD_skills, JD_exp):
#     try:
#         user_input = f"""JD: {JD_skills} and resume: {resume_skills}. Give me matching skills list from JD and resume skills"""
        
#         # Improved chat model interaction
#         chat_session  = model.start_chat(history=[
#             {"role": "user", "parts": ["jd=['devops','ml','python']"]},
#             {"role": "model", "parts": ["['devops', 'ml', 'python']"]},
#             {"role": "user", "parts": ["resume=['ml','sql','python']"]},
#             {"role": "model", "parts": ["['ml', 'python']"]},
#         ])
#         response = chat_session.send_message(user_input)
        
#         if response.text:
#             ans = response.text.strip("[]").split(", ")
#             count = len(ans)

#             # Skill score adjustment
#             skill_score = min(count / max(len(JD_skills), 1), 1.0)

#             # Improved experience score logic (using logs or thresholds)
#             resume_experience = float(resume_exp)
#             JD_exp = float(JD_exp)

#             # Example: experience score with a threshold for significance
#             experience_score = min(resume_experience / max(JD_exp, 1), 1.0)

#             # Final score with a new weight distribution (e.g., 60% skills, 40% experience)
#             final_score = (0.6 * skill_score) + (0.4 * experience_score)
#             return final_score * 100
        
#         else:
#             print("Conversation did not start successfully")
#             return 0.0
    
#     except Exception as e:
#         print(f'Error calculating score: {str(e)}')
#         return 0.0


 

@shared_task
def score_resumes_task_by_new_resume(json_data):
    # try:
    #     # Process resume data
    #     resume_skills = json_data.get('skills', [])
    #     resume_exp = json_data.get('experiance_in_number', 0.0)
    #     created_by = json_data.get('created_by', '')
    #     # Fetch all JDs from the jd_collection
    #     # print("all data",resume_skills,resume_exp,created_by)
    #     jds = list(jd_collection.find({'created_by': created_by}))
    #     client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
    #     db = client['resume_parser']
    #     score_resume_by_new_resume = db['score_resume_by_new_resume']
    #     jd_collection = db['jd_collecction']
    #     for jd in jds:
    #         try:
    #             JD_skills = jd.get('skills', [])
    #             JD_exp = jd.get('exp', 0.0)

    #             # Calculate score for the resume against the current JD
    #             score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
    #             print("score of resume",score)
    #             # Create a dictionary representing the scored resume
    #             scored_resume = {
    #                 'jd_id': str(jd['_id']),
    #                 'resume_data': json_data,  # Resume data
    #                 'jd_data': jd,  # JD data
    #                 'score': score
    #             }
    #             # Save the scored resume in the score_resume_by_new_resume collection
    #             score_resume_by_new_resume.insert_one(scored_resume)
    #         except Exception as e:
    #             logger.error(f'Error scoring resume against JDs: {str(e)}')
    #     logger.info('Scored resume against all JDs and saved successfully.')
    # except Exception as e:
    #     logger.error(f'Error scoring resume against JDs: {str(e)}')
      try:
        # Process resume data
        resume_skills = json_data.get('skills', [])
        resume_exp = json_data.get('experiance_in_number', 0.0)
        created_by = json_data.get('created_by', '')
        resume_category = json_data.get('Resume Category', '')  # Get category from resume data

        # Fetch all JDs from the jd_collection
        client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
 
        db = client['resume_parser']
        score_resume_by_new_resume = db['score_resume_by_new_resume']
        jd_collection = db['jd_collecction']

        jds = list(jd_collection.find({'created_by': created_by}))

        for jd in jds:
            try:
                JD_skills = jd.get('skills', [])
                JD_exp = jd.get('exp', 0.0)
                jd_category = jd.get('category', '')  # Get category from JD data

                # Score only if categories match
                if resume_category != jd_category:
                    print(f"Skipping JD '{jd['_id']}' due to category mismatch.")
                    continue  # Skip this JD if the category doesn't match

                # Calculate score for the resume against the current JD
                score = calculate_score(resume_skills, resume_exp, JD_skills, JD_exp)
                print("score of resume", score)

                # Create a dictionary representing the scored resume
                scored_resume = {
                    'jd_id': str(jd['_id']),
                    'resume_data': json_data,  # Resume data
                    'jd_data': jd,  # JD data
                    'score': score
                }
                
                # Save the scored resume in the score_resume_by_new_resume collection
                score_resume_by_new_resume.insert_one(scored_resume)

            except Exception as e:
                logger.error(f'Error scoring resume against JDs: {str(e)}')
                
        logger.info('Scored resume against all JDs and saved successfully.')
      except Exception as e:
        logger.error(f'Error scoring resume against JDs: {str(e)}')

def get_matching_resumes_with_scores(json_data):
    try:
        # Extract JD details
        JD_pri_skills = json_data.get('primary_skills', [])  # Expecting an array now
        JD_sec_skills = json_data.get('secondary_skills', [])  # Expecting an array now
        JD_exp = json_data.get('experience_required', 0.0)
        jd_id = json_data.get('jd_id', '')
        jd_title = json_data.get('job_title', '')
        jd_category = json_data.get('category', '')
        created_by = json_data.get('created_by', '')


        # Prepare JD details for processing
        # Prepare JD details for processing
        experience_required = int(JD_exp)
        primary_skills = [skill.strip().lower() for skill in JD_pri_skills] if isinstance(JD_pri_skills, list) else []
        secondary_skills = [skill.strip().lower() for skill in JD_sec_skills] if isinstance(JD_sec_skills, list) else []
        

        


        print(f"JD Details: Experience: {experience_required}, Primary Skills: {primary_skills}, Category: {jd_category}")

        # Connect to MongoDB
        client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
        db = client["resume_parser"]
        # resume_collection = db["resume_collection"]
        # score_resume_by_new_jd = db["score_resume_by_new_jd"]
        print("MongoDB connection successful")
        upload_record = db['upload_record']
        score_resume_by_new_jd = db['score_resume_by_new_jd']
        jd_collection = db['jd_collecction']
        resume_collection = db['resume_collection']
        
        json_data['uploaded_at'] = datetime.datetime.now()
        print("testing2....")
        
        jd_collection.insert_one(json_data)
        
        print("testing1....")
        resumes = list(resume_collection.find())
        print("testing3....")

        # Fetch resumes matching the experience and category
        initial_query = {
            "experiance_in_number": {"$gte": experience_required},
            "Resume_Category": jd_category
        }
        print("Query to fetch resumes:", initial_query)
        potential_resumes = list(resume_collection.find(initial_query))
        print(f"Potential resumes fetched: {len(potential_resumes)}")

        # Filter resumes by skill threshold
        matching_resumes = []
        for resume in potential_resumes:
            resume_skills = {skill.strip().lower() for skill in resume.get("skills", [])}
            matching_skill_count = len(resume_skills.intersection(primary_skills))
            required_skill_threshold = max(1, int(0.1 * len(primary_skills)))  # At least one skill match

            if matching_skill_count >= required_skill_threshold:
                matching_resumes.append(resume)

        print(f"Matching resumes after skill filtering: {len(matching_resumes)}")

        # Calculate scores and save results
        scored_resumes = []
        for resume in matching_resumes:
            score = calculate_score(resume, {
                "experience_required": experience_required,
                "primary_skills": JD_pri_skills,
                "secondary_skills": JD_sec_skills,
                "job_role": json_data.get("job_role", ""),
                "domain_specific": json_data.get("domain_specific", [])
            })

            # scored_resume = {
            #     "jd_data": {
            #         "jd_id": str(jd_id),
            #         "jd_title": jd_title,
            #         "jd_category": jd_category
            #     },
            #     "resume_data": resume,
            #     "created_by": created_by
            # }
            scored_resume = {
                    'resume_id': str(resume['_id']),
                    'resume_data': resume,
                    'score': score,
                    'jd_data': json_data,
                    'created_by': created_by
                }
            scored_resumes.append(scored_resume)

        # Save to MongoDB
        if scored_resumes:
            score_resume_by_new_jd.insert_many(scored_resumes)
            print(f"Scores saved for {len(scored_resumes)} resumes.")
        else:
            print("No scores to save. No matching resumes found.")

    except Exception as e:
        print(f"Error in processing resumes: {e}")



def file_paths_and_convert_docx(data_obj):
    try:
        resume_data = data_obj["resume_data"]
        jd_data = data_obj["jd_data"]
        que_level = data_obj.get("que_level")
        len_que = data_obj.get("len_que")
        created_by = data_obj["created_by"]
        ref_no = data_obj["ref_num"]
 
        resume_dic = resume_data
        jd_dic = jd_data
       
        if len(resume_dic) > 0 and len(jd_dic) > 0:
            exp_str = jd_dic.get("experience_required", "0")  # Ensure correct key
            pattern = r'\b\d+(\.\d+)?\b'
            match = re.search(pattern, exp_str)
            exp = float(match.group()) if match else 0
            if exp < 2:
                prompt = "focus  on generating 70% technical questions and 30% project-based questions"
            elif 2 <= exp <= 5:
                prompt = "focus on generating 50% technical questions and 50% project-based questions"
            else:
                prompt = "focus on generating 80% project-based and situational questions and 20% technical questions"
 
 
            ## TO Identify the Top5 Skills
            Instruction_top5skills = """the skills format to follow is below
            ["skill1","skill2","skill3","skill4","skill5"]
 
            """
 
            top5skills_input_text = f"""Analyze the provided job description-{jd_dic} and extract the top 5 primary skills most relevant to the role.
            Focus solely on the skills explicitly mentioned or emphasized in the job description, ensuring that the selection prioritizes those that are critical for the role and essential for its success.
            Remember, in the provided job description, majorly concentrate on the primary skills mentioned and aslo consider from the Secondary skills.
            In the output should always in the string format and just give me the Skills not the description for them. and alos dont include any introduction lines.
            """
 
 
            agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                        instructions = Instruction_top5skills,
                        add_references_to_prompt=True)
 
            # agent.print_response("give me the name from the uploaded file")
            run: RunResponse = agent.run(top5skills_input_text)
            top5skills = run.content
            print(top5skills)

            Instruction_name = """
            name : Candidate name 
                """
            name_input_text = f"""take the resume information from the {resume_dic} and extract the candidate name mentioned in the resume.

            """

            agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                        instructions = Instruction_name,
                        add_references_to_prompt=True)
 
            # agent.print_response("give me the name from the uploaded file")
            run: RunResponse = agent.run(name_input_text)
            candidate_name = run.content
            print(candidate_name)
 
            Instruction_quesn = """
            Strictly adhere the Below is format generate the Questions.
            "interview_questions": [
                {
                "question": "Please introduce yourself."
                },
                {
                "question": "Discuss one of your projects in detail."
                },
                {
                "question": ""
                },
                .
                .
                .
            ]
            """
            # Extract the candidate's name using regex
            # match = re.search(r"'name'\s*:\s*'([^']+)'", resume_dic)
            # candidate_name = match.group(1) if match else "Candidate"
            # candidate_name ="Tarun"
            language_selected = "english"
            

 
            input_text_quesn = f"""
'Resume Information': {resume_dic} and 'Job Description Information': {jd_dic},
Analyze the candidate's skills, qualifications, experience, and projects from both the resume and the job description.
Generate {len_que} interview questions at {que_level} level that comprehensively cover the following:
1. The key technical skills and competencies listed in the resume and job description.
2. The candidate's relevant experience, including specific achievements and challenges.
3. The qualifications and certifications mentioned in the resume, ensuring relevance to the role.
4. The projects and roles the candidate has worked on, focusing on problem-solving, results, and impact.
5. The alignment between the candidate's background and the requirements of the job description.

In addition to these {len_que} questions, include the following two mandatory questions:
1. Start with a personalized greeting using the candidate’s name extracted from the resume. Welcome them warmly and thank them for attending the interview with the AI Interview Assistant. Then, ask them to introduce themselves in a natural and engaging manner.
   Example: "Hello {candidate_name}, welcome! Thank you for joining us today for this interview powered by the AI Interview Assistant. We’re excited to learn more about you. To start, could you please introduce yourself?"
   
2. Choose one of the projects mentioned in the resume and ask the candidate to explain it in detail, focusing on their role, approach, and outcomes.
   Remember, don't give the choice to select the projects.

Incorporate {top5skills} into the question generation process:
- 75% of the questions should focus on these top 5 skills listed in the job description.
- The remaining 25% should cover other important concepts and requirements mentioned in the job description.

Additionally, apply the logic defined in {prompt} for question distribution based on the candidate's experience level (`exp`).

Ensure the questions are balanced between technical and behavioral aspects, and include follow-up questions where necessary.
the output must be in the English language only. 
Strictly adhere to the the language mentioned - English and question level {que_level}. Store all the questions in JSON format."""

 
            agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                        instructions = Instruction_quesn,
                        add_references_to_prompt=True)
 
            # agent.print_response("give me the name from the uploaded file")
            run: RunResponse = agent.run(input_text_quesn)
            # questions = (run.content).strip().strip("```").replace("text", "").strip()
            questions = (run.content).strip().strip("```json").strip()
            print(questions)
 
           # Parse the JSON string into a Python dictionary
            questions_dict = json.loads(questions)
 
            # Extract questions from the list and store them in `questions_list_str`
            questions_list_str = [q["question"] for q in questions_dict["interview_questions"]]
 
 
            if len(questions_list_str) > 0:
                print("questions_list_str is not empty.")
            else:
                print("questions_list_str is empty.")
           
            questions = []
            for question_str in questions_list_str:
                print('came in question', question_str)
                
                # Ensure the question_str is formatted properly as a JSON string
                question_str = question_str.strip()
                
                # Add proper quotes around the "question" key and value if they aren't already present
                if not question_str.startswith('"question":'):
                    question_str = '"question": "' + question_str + '"'
                
                # Wrap the question_str with curly braces to make it a valid JSON object
                question_json_str = '{' + question_str + '}'
                
                try:
                    # Load the string into a JSON object
                    question_json = json.loads(question_json_str)
                    question_json["answer"] = ""  # Add the "answer" key with an empty string
                    
                    # Append the question to the questions list
                    questions.append(question_json)
                except json.JSONDecodeError as e:
                    print(f"Error decoding JSON: {e}")
                    continue  # Skip invalid question_str
 
            # Save to MongoDB
            client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@192.168.0.16:27017/")
 
            db = client["resume_parser"]
            collection = db["interview_collection"]
 
            upload_data = {
                "resume_file_name": resume_dic.get("file_name"),
                "reference_number": ref_no,
                "resume_data": resume_dic,
                "input_type": data_obj.get("input_type"),
                "que_level": que_level,
                "len_que": len_que,
                "uploaded_by": created_by,
                "original_resume_file_name": resume_dic.get("original_file_name"),
                "jd_data": jd_dic,
                "questions": questions,
                "interviewType": "AI"
            }
 
            result = collection.insert_one(upload_data)
            print("testing...")
            print("this is result", result)
            object_id = str(result.inserted_id)  # Get the ObjectId of the inserted document
            print("This is the objectid",object_id)
            client.close()
 
            # Extract email from resume_data
            # email = resume_dic.get("email")
            email = re.sub(r"\s+", "", resume_dic.get("email")) if resume_dic.get("email") else ""
            desig = jd_dic.get("job_title")
            candidatename = resume_dic.get("name")
            print("This is the objectid",email)
            print("This is the designation",desig)
            print("This is the designation",candidatename)
 
 
            if email:
                api_payload = {
                    "email": email,
                    "userInterviewId": object_id,
                    "interviewProfile": desig,
                    "candidateName": candidatename
                }
 
                # Call the Java API
                api_url = "http://localhost:8083/interview/sendemail"  # Replace with your actual API URL
                response = requests.post(api_url, json=api_payload ,  verify=False)

                if response.status_code == 200:
                    print("API call was successful.")
                else:
                    print(f"API call failed with status code: {response.status_code}")
            else:
                print("Email not found in resume_data.")

            client.close()
        else:
            print(f'Resume or JD data is empty')
    except Exception as e:
        print(f'Error in processing data here:' + str(e))


from pymongo import MongoClient
import json
import requests
import re
from celery import shared_task

def coding_question_generation(data_obj):
    try:
        # Extracting data from the request object
        resume_data = data_obj["resume_data"]
        jd_data = data_obj["jd_data"]
        que_level = data_obj.get("que_level")
        len_que = data_obj.get("len_que")
        created_by = data_obj["created_by"]
        ref_no = data_obj["ref_num"]
        programming_language = data_obj.get("programming_language")  # Dynamic input
        num_question = data_obj.get("num_question")  # Removed default value
        
        if num_question is None:
            raise ValueError("Number of questions is required")

        qsn_level = data_obj.get("qsn_level", "Medium")
        particular_skill = data_obj.get("particular_skill", "")

        resume_dic = resume_data
        jd_dic = jd_data

        # Instructions for question generation
        Instruction_qn_gen = """
            Generate clear, precise, and role-specific technical coding questions tailored to the provided Job Description (JD) details.
            Ensure the questions assess both fundamental and advanced skills relevant to the role. 
            Each question should align with real-world scenarios to evaluate practical problem-solving abilities.

            Ensure the questions include:
            - Title
            - Problem Statement
            - Input Format
            - Output Format
            
            Each question should target specific skills listed in the job description.
            Provide exactly the number of questions specified.
            Ensure JSON format in output.
        """

        input_text = f"""
            Objective:
            Generate coding assessment questions tailored to the candidate’s profile and the job description.

            Instructions:
            Resume Information: {resume_data if resume_data else "Not Provided"}
            Job Description: {jd_data if jd_data else "Not Provided"}
            Experience Level: {que_level if que_level else "Not Provided"}
            Programming Language: {programming_language}
            Number of Questions: {num_question}
            Question Level: {qsn_level}
            Skill-Specific: {particular_skill if particular_skill else "Not Provided"}

            Important: 
            - Provide all questions in JSON format.
            - Include clear input and output format for each question.
        """

        agent = Agent(model=Gemini(id="gemini-2.0-flash-exp"),
                      instructions=Instruction_qn_gen,
                      add_references_to_prompt=True)

        run = agent.run(input_text)
        questions_str = run.content.strip().strip("```json").strip()

        print(f"Raw Model Output:\n{questions_str}")

        if not questions_str:
            print("No response from the model.")
            return []

        try:
            questions_data = json.loads(questions_str)
            questions_list = questions_data.get("questions", [])
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON from model: {e}")
            return []

        # Extract and clean questions directly
        questions = []
        for q in questions_list:
            question = {
                "title": q.get("title", ""),
                "problem_statement": q.get("problem_statement", ""),
                "input_format": q.get("input_format", ""),
                "output_format": q.get("output_format", ""),
                "skills_tested": q.get("skills_tested", []),
                "answer": ""
            }
            questions.append(question)

        if questions:
            print(f"Extracted questions:\n{json.dumps(questions, indent=4)}")
        else:
            print("No questions extracted.")

        # Save to MongoDB
        client = MongoClient(f"mongodb://{mongo_username}:{mongo_password}@localhost:27017/")
        db = client["resume_parser"]
        collection = db["interview_collection"]

        upload_data = {
            "resume_file_name": resume_dic.get("file_name"),
            "reference_number": ref_no,
            "resume_data": resume_dic,
            "input_type": data_obj.get("input_type"),
            "que_level": que_level,
            "len_que": len_que,
            "uploaded_by": created_by,
            "original_resume_file_name": resume_dic.get("original_file_name"),
            "jd_data": jd_dic,
            "questions": questions,
            "interview_type": "coding"
        }

        result = collection.insert_one(upload_data)
        object_id = str(result.inserted_id)
        print(f"Inserted document ID: {object_id}")
        client.close()

        # Send Email if available
        email = re.sub(r"\s+", "", resume_dic.get("email", ""))
        desig = jd_dic.get("job_title", "")
        candidate_name = resume_dic.get("name", "")

        if email:
            api_payload = {
                "email": email,
                "userInterviewId": object_id,
                "interviewProfile": desig,
                "candidateName": candidate_name
            }

            api_url = "https://aippoint.ai/aippoint-spring-codeassessment/interview/sendemail"
            response = requests.post(api_url, json=api_payload, verify=False)

            if response.status_code == 200:
                print("Email sent successfully.")
            else:
                print(f"Failed to send email. Status code: {response.status_code}")

    except Exception as e:
        print(f"Error in coding_question_generation: {e}")

# Celery task
@shared_task
def celery_coding_question(file_data):
    try:
        for data in file_data:
            coding_question_generation(data)
        return True
    except Exception as e:
        print(f"Error in celery task: {e}")
        return False
    

  

@shared_task
def process_videos_task(video_urls, meeting_id, object_id):
    """
    Celery task to handle video merging, uploading, fetching, and proctoring API calls.
    """
    temp_dir = tempfile.mkdtemp()
    try:
        video_files = []
        for idx, url in enumerate(video_urls):
            if not url.startswith("http"):
                logger.warning(f"Skipping invalid URL: {url}")
                continue

            video_path = os.path.join(temp_dir, f"video_{idx}.mp4")
            try:
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                with open(video_path, "wb") as f:
                    f.write(response.content)
                video_files.append(video_path)
                logger.info(f"Downloaded video {idx + 1} from {url}")
            except requests.RequestException as e:
                logger.error(f"Failed to download video from {url}: {e}")
                continue

        if video_files:
            logger.info(f"Merging {len(video_files)} videos")
            clips = [VideoFileClip(vid) for vid in video_files]
            final_clip = concatenate_videoclips(clips, method="compose")
            output_path = os.path.join(temp_dir, "merged_video.mp4")
            final_clip.write_videofile(output_path, codec="libx264", fps=24, preset='ultrafast')
            logger.info(f"Merged video saved at {output_path}")

            for clip in clips:
                clip.close()
            final_clip.close()

            with open(output_path, 'rb') as file:
                files = {'file': ('merged_video.mp4', file, 'video/mp4')}
                data = {'meetingId': meeting_id}
                upload_response = requests.put(f"{SPRING_BOOT_URL}upload", files=files, data=data, timeout=10, verify=False)
                upload_response.raise_for_status()

                if upload_response.status_code == 200:
                    logger.info("Video uploaded successfully to backend")
                    merged_video_data = fetch_merged_video_from_backend(meeting_id)
                    if merged_video_data:
                        call_proctoring_apis(merged_video_data, object_id)
                        logger.info("Proctoring APIs called successfully")
                    else:
                        logger.error("Failed to fetch merged video from backend")
                else:
                    logger.error(f"Failed to upload video: {upload_response.text}")
        else:
            logger.error("No valid videos were downloaded")
    except Exception as e:
        logger.exception(f"An error occurred in task: {e}")
    finally:
        if os.path.exists(temp_dir):
            max_attempts = 5
            attempt = 0
            while attempt < max_attempts:
                try:
                    shutil.rmtree(temp_dir, ignore_errors=False)
                    logger.info("Cleaned up temp directory")
                    break
                except PermissionError as e:
                    attempt += 1
                    logger.warning(f"PermissionError during cleanup, attempt {attempt}/{max_attempts}: {e}")
                    time.sleep(1)
                except Exception as e:
                    logger.error(f"Unexpected error during cleanup: {e}")
                    break
            if attempt == max_attempts:
                logger.error(f"Failed to clean up temp directory {temp_dir} after {max_attempts} attempts")

def fetch_merged_video_from_backend(meeting_id, max_retries=3, retry_delay=5):
    """
    Fetches the merged video file from the backend for the given meeting ID as binary data (blob).
    """
    try:
        print("This is meeting_id", meeting_id)
        attempt = 0
        while attempt < max_retries:
            try:
                response = requests.get(f"{SPRING_BOOT_URL}video/{meeting_id}", stream=True, timeout=10, verify=False)
                print("Came to fetch merged video")
                response.raise_for_status()
                video_data = io.BytesIO(response.content)
                logger.info(f"Fetched merged video for meeting_id: {meeting_id} as binary data")
                return video_data
            except requests.RequestException as e:
                if response.status_code == 404:
                    logger.warning(f"Video not found for meeting_id {meeting_id}, attempt {attempt + 1}/{max_retries}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        attempt += 1
                        continue
                    logger.error(f"Failed to fetch merged video after {max_retries} attempts: {e}")
                    return None
                else:
                    logger.error(f"Failed to fetch merged video: {e}")
                    return None
        return None
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        return None

def call_proctoring_apis(video_data, object_id, max_retries=3, retry_delay=8, timeout=300):
    """
    Sends the in-memory video data (blob) to various proctoring APIs with retries and increased timeout.
    """
    if not video_data:
        logger.error("No video data provided, skipping proctoring.")
        return

    endpoints = [
        "eye-tracking-analysis/",
        "voice_analysis/",
        "noise_analysis/",
        "multiple-faces-detection/",
    ]

    video_bytes = video_data.getvalue()

    for endpoint in endpoints:
        attempt = 0
        while attempt < max_retries:
            try:
                video_buffer = io.BytesIO(video_bytes)
                files = {"video_file": ("merged_video.mp4", video_buffer, "video/mp4")}
                logger.info(f"Calling proctoring API: {endpoint}")
                response = requests.post(f"{PROCTORING_API_BASE_URL}{endpoint}", files=files, timeout=timeout, verify=False)
                response.raise_for_status()
                analysis_result = response.json()
                if endpoint == "eye-tracking-analysis/":
                    save_proctoring_results("save_eye_tracking_results/", analysis_result, object_id)
                elif endpoint == "noise_analysis/":
                    save_proctoring_results("save_noise_detection_results/", analysis_result, object_id)
                elif endpoint == "voice_analysis/":
                    save_proctoring_results("save_voice_analysis_results/", analysis_result, object_id)
                elif endpoint == "multiple-faces-detection/":
                    save_proctoring_results("save_multiple_face_analysis/", analysis_result, object_id)
                else:
                    logger.warning(f"No save API mapping for endpoint {endpoint}")
                logger.info(f"Proctoring API {endpoint} response: {analysis_result}")
                


                break
            except requests.Timeout as e:
                attempt += 1
                logger.warning(f"Timeout calling {endpoint}, attempt {attempt}/{max_retries}: {e}")
                if attempt < max_retries:
                    time.sleep(retry_delay)
                    continue
                logger.error(f"Failed to call {endpoint} after {max_retries} attempts: {e}")
            except requests.RequestException as e:
                logger.error(f"Failed to call {endpoint}: {e}")
                break
            finally:
                video_buffer.close()


def save_proctoring_results(api_endpoint, analysis_result, object_id):
    try:
        save_response = requests.post(
            f"https://aippoint.ai/aippoint-django-service/{api_endpoint}",
            json={
                "object_id": str(object_id),
                "analysis_result": analysis_result,
            },
            headers={"Content-Type": "application/json"},
            verify=False)

        print("came to save analysis")
        save_response.raise_for_status()
        logger.info(f"Saved results to {api_endpoint} for ObjectId: {object_id}")
    except Exception as e:
        logger.exception(f"Failed to save results to {api_endpoint} for ObjectId {object_id}: {e}")

